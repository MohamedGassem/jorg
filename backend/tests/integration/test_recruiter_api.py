# backend/tests/integration/test_recruiter_api.py
import io
from typing import Any

from docx import Document  # type: ignore[import-untyped,unused-ignore]
from httpx import AsyncClient

# ---- Auth & role guards -----------------------------------------------------


async def test_get_recruiter_profile_requires_auth(client: AsyncClient) -> None:
    r = await client.get("/recruiters/me/profile")
    assert r.status_code == 401


async def test_candidate_cannot_get_recruiter_profile(
    client: AsyncClient, candidate_headers: dict[str, str]
) -> None:
    r = await client.get("/recruiters/me/profile", headers=candidate_headers)
    assert r.status_code == 403


# ---- RecruiterProfile -------------------------------------------------------


async def test_get_recruiter_profile_auto_creates(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    r = await client.get("/recruiters/me/profile", headers=recruiter_headers)
    assert r.status_code == 200
    data = r.json()
    assert data["first_name"] is None
    assert data["organization_id"] is None
    assert "id" in data


async def test_update_recruiter_profile(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    r = await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"first_name": "Bob", "last_name": "Smith", "job_title": "Talent Manager"},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["first_name"] == "Bob"
    assert data["job_title"] == "Talent Manager"


# ---- Organization -----------------------------------------------------------


async def test_create_organization(client: AsyncClient, recruiter_headers: dict[str, str]) -> None:
    r = await client.post(
        "/organizations",
        headers=recruiter_headers,
        json={"name": "Acme Corp"},
    )
    assert r.status_code == 201
    data = r.json()
    assert data["name"] == "Acme Corp"
    assert data["slug"] == "acme-corp"
    assert "id" in data


async def test_create_organization_slug_is_unique(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    await client.post("/organizations", headers=recruiter_headers, json={"name": "Dupont SA"})
    r2 = await client.post("/organizations", headers=recruiter_headers, json={"name": "Dupont SA"})
    assert r2.status_code == 201
    assert r2.json()["slug"] == "dupont-sa-1"


async def test_get_organization(client: AsyncClient, recruiter_headers: dict[str, str]) -> None:
    create = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "Test Inc"}
    )
    org_id = create.json()["id"]
    r = await client.get(f"/organizations/{org_id}", headers=recruiter_headers)
    assert r.status_code == 200
    assert r.json()["name"] == "Test Inc"


async def test_get_organization_not_found(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    r = await client.get(
        "/organizations/00000000-0000-0000-0000-000000000000",
        headers=recruiter_headers,
    )
    assert r.status_code == 404


async def test_recruiter_can_link_to_organization(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    org = await client.post("/organizations", headers=recruiter_headers, json={"name": "My Firm"})
    org_id = org.json()["id"]
    r = await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )
    assert r.status_code == 200
    assert r.json()["organization_id"] == org_id


async def test_candidate_cannot_create_organization(
    client: AsyncClient, candidate_headers: dict[str, str]
) -> None:
    r = await client.post(
        "/organizations",
        headers=candidate_headers,
        json={"name": "Should Fail"},
    )
    assert r.status_code == 403


# ---- Template upload --------------------------------------------------------


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Create a minimal .docx in memory and return its bytes."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _setup_org_and_link(client: AsyncClient, recruiter_headers: dict[str, str]) -> str:
    """Helper: create an org and link the recruiter to it. Returns org_id."""
    org = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "Template Corp"}
    )
    org_id: str = org.json()["id"]
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )
    return org_id


async def test_upload_template_detects_placeholders(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    org_id = await _setup_org_and_link(client, recruiter_headers)
    docx_bytes = _make_docx_bytes(["Nom: {{NOM}}", "Prénom: {{PRENOM}}", "Titre: {{TITRE}}"])
    r = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "Mon Template"},
        files={
            "file": (
                "template.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 201
    data = r.json()
    assert data["name"] == "Mon Template"
    assert "{{NOM}}" in data["detected_placeholders"]
    assert "{{PRENOM}}" in data["detected_placeholders"]
    assert "{{TITRE}}" in data["detected_placeholders"]
    # Le rendu mock reussit (balises inconnues rendues vides) : valide avec avertissements.
    assert data["is_valid"] is True
    assert set(data["unknown_placeholders"]) == {"{{NOM}}", "{{PRENOM}}", "{{TITRE}}"}


async def test_list_templates(client: AsyncClient, recruiter_headers: dict[str, str]) -> None:
    org_id = await _setup_org_and_link(client, recruiter_headers)
    docx_bytes = _make_docx_bytes(["{{NOM}}"])
    for name in ["T1", "T2"]:
        await client.post(
            f"/organizations/{org_id}/templates",
            headers=recruiter_headers,
            data={"name": name},
            files={
                "file": (
                    "t.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    r = await client.get(f"/organizations/{org_id}/templates", headers=recruiter_headers)
    assert r.status_code == 200
    assert len(r.json()) == 2


async def test_upload_template_with_standard_placeholders_is_valid(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    org_id = await _setup_org_and_link(client, recruiter_headers)
    docx_bytes = _make_docx_bytes(["{{last_name}} {{first_name}}"])
    r = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "T"},
        files={
            "file": (
                "t.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 201
    assert r.json()["is_valid"] is True
    assert r.json()["unknown_placeholders"] == []


async def test_delete_template(client: AsyncClient, recruiter_headers: dict[str, str]) -> None:
    org_id = await _setup_org_and_link(client, recruiter_headers)
    docx_bytes = _make_docx_bytes(["{{NOM}}"])
    upload = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "ToDelete"},
        files={
            "file": (
                "t.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    template_id = upload.json()["id"]

    r = await client.delete(
        f"/organizations/{org_id}/templates/{template_id}",
        headers=recruiter_headers,
    )
    assert r.status_code == 204

    list_r = await client.get(f"/organizations/{org_id}/templates", headers=recruiter_headers)
    assert len(list_r.json()) == 0


async def test_recruiter_cannot_access_other_org_templates(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    """A recruiter not linked to an org gets 403 on its templates."""
    # Create a second recruiter that will own the org
    await client.post(
        "/auth/register",
        json={"email": "owner_recruiter@test.com", "password": "pass1234", "role": "recruiter"},
    )
    owner_login = await client.post(
        "/auth/login", json={"email": "owner_recruiter@test.com", "password": "pass1234"}
    )
    owner_headers = {"Authorization": f"Bearer {owner_login.json()['access_token']}"}

    # Owner creates the org (auto-linked as creator)
    org = await client.post("/organizations", headers=owner_headers, json={"name": "Other Corp"})
    org_id = org.json()["id"]
    # Original recruiter is not linked to this org
    r = await client.get(f"/organizations/{org_id}/templates", headers=recruiter_headers)
    assert r.status_code == 403


async def test_upload_template_excludes_control_syntax_from_detected(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
) -> None:
    """Uploading a docxtpl template must surface only profile-field placeholders.

    Jinja2 block tags ({%p for ... %}) and loop variable access ({{exp.*}},
    {{sk.*}}) must not appear in detected_placeholders — only standalone
    profile fields like {{last_name}} should be listed.
    """
    org_id = await _setup_org_and_link(client, recruiter_headers)
    docx_bytes = _make_docx_bytes(
        [
            "Candidat : {{last_name}} {{first_name}}",
            "{%p for exp in experiences %}",
            "{{exp.client_name}} — {{exp.role}}",
            "{%p endfor %}",
        ]
    )

    r = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "With block"},
        files={
            "file": (
                "with_block.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert r.status_code == 201, r.text
    body = r.json()
    assert "{{last_name}}" in body["detected_placeholders"]
    assert "{{first_name}}" in body["detected_placeholders"]
    # Loop variables and block tags must be excluded
    assert not any(p.startswith("{{exp.") for p in body["detected_placeholders"])
    assert not any("{%p" in p for p in body["detected_placeholders"])
    assert (
        body["is_valid"] is True
    )  # all detected placeholders are known fields — auto-mapped on upload


# ---- Accessible candidates --------------------------------------------------


async def test_list_accessible_candidates_empty(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    org = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "Access Corp"}
    )
    org_id = org.json()["id"]
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )
    r = await client.get(f"/organizations/{org_id}/candidates", headers=recruiter_headers)
    assert r.status_code == 200
    assert r.json() == []


async def test_list_accessible_candidates_returns_granted(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    org = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "Access Corp"}
    )
    org_id = org.json()["id"]
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )
    await client.put(
        "/candidates/me/profile",
        headers=candidate_headers,
        json={"first_name": "Alice", "last_name": "Dupont"},
    )
    inv = await client.post(
        f"/organizations/{org_id}/invitations",
        headers=recruiter_headers,
        json={"candidate_email": "candidate@test.com"},
    )
    token = inv.json()["token"]
    r = await client.post(f"/invitations/{token}/accept", headers=candidate_headers)
    assert r.status_code == 201

    r = await client.get(f"/organizations/{org_id}/candidates", headers=recruiter_headers)
    assert r.status_code == 200
    data = r.json()
    assert len(data) == 1
    assert data[0]["email"] == "candidate@test.com"
    assert data[0]["first_name"] == "Alice"
    assert data[0]["last_name"] == "Dupont"
    assert "user_id" in data[0]


async def test_list_accessible_candidates_excludes_revoked(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    org = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "Access Corp"}
    )
    org_id = org.json()["id"]
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )
    inv = await client.post(
        f"/organizations/{org_id}/invitations",
        headers=recruiter_headers,
        json={"candidate_email": "candidate@test.com"},
    )
    accepted = await client.post(
        f"/invitations/{inv.json()['token']}/accept", headers=candidate_headers
    )
    grant_id = accepted.json()["id"]
    await client.delete(f"/access/me/{grant_id}", headers=candidate_headers)

    r = await client.get(f"/organizations/{org_id}/candidates", headers=recruiter_headers)
    assert r.status_code == 200
    assert r.json() == []


async def test_list_accessible_candidates_requires_membership(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    # Create a second recruiter that owns the org
    await client.post(
        "/auth/register",
        json={"email": "owner2@test.com", "password": "pass1234", "role": "recruiter"},
    )
    owner_login = await client.post(
        "/auth/login", json={"email": "owner2@test.com", "password": "pass1234"}
    )
    owner_headers = {"Authorization": f"Bearer {owner_login.json()['access_token']}"}

    org = await client.post("/organizations", headers=owner_headers, json={"name": "Other Org"})
    org_id = org.json()["id"]
    # Original recruiter is not linked to this org
    r = await client.get(f"/organizations/{org_id}/candidates", headers=recruiter_headers)
    assert r.status_code == 403


async def test_list_accessible_candidates_forbids_candidate_role(
    client: AsyncClient, candidate_headers: dict[str, str]
) -> None:
    r = await client.get(
        "/organizations/00000000-0000-0000-0000-000000000000/candidates",
        headers=candidate_headers,
    )
    assert r.status_code == 403


# ---- Template file download -------------------------------------------------


async def test_preview_template_renders_mock_docx(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    org_id = await _setup_org_and_link(client, recruiter_headers)
    docx_bytes = _make_docx_bytes(["{{first_name}} {{last_name}}"])
    up = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "T"},
        files={
            "file": (
                "t.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    template_id = up.json()["id"]
    r = await client.get(
        f"/organizations/{org_id}/templates/{template_id}/preview",
        headers=recruiter_headers,
    )
    assert r.status_code == 200
    assert r.content[:2] == b"PK"
    assert "attachment" in r.headers["content-disposition"]


async def test_templatize_returns_503_when_llm_disabled(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    org_id = await _setup_org_and_link(client, recruiter_headers)
    docx_bytes = _make_docx_bytes(["Jean Dupont"])
    up = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "T"},
        files={
            "file": (
                "t.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    template_id = up.json()["id"]
    r = await client.post(
        f"/organizations/{org_id}/templates/{template_id}/templatize",
        headers=recruiter_headers,
    )
    assert r.status_code == 503


async def test_capabilities_reports_assisted_templating_flag(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    r = await client.get("/templates/capabilities", headers=recruiter_headers)
    assert r.status_code == 200
    assert r.json() == {"assisted_templating": False}


async def test_activate_draft_template(
    client: AsyncClient, recruiter_headers: dict[str, str], monkeypatch: Any
) -> None:
    from services.documents.templatize_service import TemplatizeOutcome

    async def fake_pipeline(client_: Any, model: str, path: str) -> TemplatizeOutcome:
        doc = Document()
        doc.add_paragraph("{{first_name}}")
        buf = io.BytesIO()
        doc.save(buf)
        return TemplatizeOutcome(
            docx_bytes=buf.getvalue(),
            report={"mappings": [], "warnings": [], "rejected": [], "render_error": None},
            render_error=None,
        )

    monkeypatch.setattr(
        "api.routes.organizations.templatize_service.run_templatize_pipeline", fake_pipeline
    )
    monkeypatch.setattr(
        "api.routes.organizations.llm_client.get_anthropic_client", lambda: object()
    )

    org_id = await _setup_org_and_link(client, recruiter_headers)
    docx_bytes = _make_docx_bytes(["Jean Dupont"])
    up = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "T"},
        files={
            "file": (
                "t.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    template_id = up.json()["id"]

    r = await client.post(
        f"/organizations/{org_id}/templates/{template_id}/templatize",
        headers=recruiter_headers,
    )
    assert r.status_code == 200
    assert r.json()["status"] == "draft"
    assert r.json()["templatize_report"]["render_error"] is None

    r = await client.post(
        f"/organizations/{org_id}/templates/{template_id}/activate",
        headers=recruiter_headers,
    )
    assert r.status_code == 200
    assert r.json()["status"] == "active"


async def test_templatize_always_runs_from_source(
    client: AsyncClient, recruiter_headers: dict[str, str], monkeypatch: Any
) -> None:
    # Regression: re-templatizing must re-run from the original upload, not from
    # the previous templatized draft, and must drop the superseded draft file.
    from pathlib import Path

    from services.documents.templatize_service import TemplatizeOutcome

    seen_paths: list[str] = []

    async def fake_pipeline(client_: Any, model: str, path: str) -> TemplatizeOutcome:
        seen_paths.append(path)
        doc = Document()
        doc.add_paragraph("{{first_name}}")
        buf = io.BytesIO()
        doc.save(buf)
        return TemplatizeOutcome(
            docx_bytes=buf.getvalue(),
            report={"mappings": [], "warnings": [], "rejected": [], "render_error": None},
            render_error=None,
        )

    monkeypatch.setattr(
        "api.routes.organizations.templatize_service.run_templatize_pipeline", fake_pipeline
    )
    monkeypatch.setattr(
        "api.routes.organizations.llm_client.get_anthropic_client", lambda: object()
    )

    org_id = await _setup_org_and_link(client, recruiter_headers)
    up = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "T"},
        files={
            "file": (
                "t.docx",
                _make_docx_bytes(["Jean Dupont"]),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    template_id = up.json()["id"]

    await client.post(
        f"/organizations/{org_id}/templates/{template_id}/templatize",
        headers=recruiter_headers,
    )
    await client.post(
        f"/organizations/{org_id}/templates/{template_id}/templatize",
        headers=recruiter_headers,
    )

    # Both runs received the same (source) path, not the first draft's output.
    assert seen_paths[0] == seen_paths[1]
    # The superseded draft from run 1 was deleted; the source path still exists.
    assert Path(seen_paths[0]).exists()


async def test_download_template_file_ok(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    """Recruiter can download the .docx file of a template they own."""
    org_r = await client.post("/organizations", json={"name": "DL Org"}, headers=recruiter_headers)
    org_id = org_r.json()["id"]
    await client.put(
        "/recruiters/me/profile", json={"organization_id": org_id}, headers=recruiter_headers
    )

    import io

    from docx import Document

    buf = io.BytesIO()
    Document().save(buf)
    buf.seek(0)
    upload_r = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        files={
            "file": (
                "test.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"name": "Test Template"},
    )
    assert upload_r.status_code == 201
    template_id = upload_r.json()["id"]

    r = await client.get(
        f"/organizations/{org_id}/templates/{template_id}/file",
        headers=recruiter_headers,
    )
    assert r.status_code == 200
    assert "application/vnd.openxmlformats" in r.headers["content-type"]
    assert "attachment" in r.headers.get("content-disposition", "")


async def test_download_template_file_wrong_org_returns_403(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    """A recruiter cannot download a template from another organization."""
    org_a = await client.post("/organizations", json={"name": "Org A"}, headers=recruiter_headers)
    org_a_id = org_a.json()["id"]
    await client.put(
        "/recruiters/me/profile", json={"organization_id": org_a_id}, headers=recruiter_headers
    )

    await client.post(
        "/auth/register",
        json={"email": "rec2@test.com", "password": "pass1234", "role": "recruiter"},
    )
    login2 = await client.post(
        "/auth/login", json={"email": "rec2@test.com", "password": "pass1234"}
    )
    headers2 = {"Authorization": f"Bearer {login2.json()['access_token']}"}

    org_b = await client.post("/organizations", json={"name": "Org B"}, headers=headers2)
    org_b_id = org_b.json()["id"]
    await client.put("/recruiters/me/profile", json={"organization_id": org_b_id}, headers=headers2)

    import io

    from docx import Document

    buf = io.BytesIO()
    Document().save(buf)
    buf.seek(0)
    upload_r = await client.post(
        f"/organizations/{org_b_id}/templates",
        headers=headers2,
        files={
            "file": (
                "test.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"name": "Org B Template"},
    )
    template_id = upload_r.json()["id"]

    r = await client.get(
        f"/organizations/{org_b_id}/templates/{template_id}/file",
        headers=recruiter_headers,
    )
    assert r.status_code == 403


async def test_download_template_file_not_found_returns_404(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    org_r = await client.post("/organizations", json={"name": "NF Org"}, headers=recruiter_headers)
    org_id = org_r.json()["id"]
    await client.put(
        "/recruiters/me/profile", json={"organization_id": org_id}, headers=recruiter_headers
    )

    import uuid

    r = await client.get(
        f"/organizations/{org_id}/templates/{uuid.uuid4()}/file",
        headers=recruiter_headers,
    )
    assert r.status_code == 404


# ---- join_code --------------------------------------------------------------


async def test_create_organization_has_join_code(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    r = await client.post(
        "/organizations",
        headers=recruiter_headers,
        json={"name": "JoinCode Corp"},
    )
    assert r.status_code == 201
    body = r.json()
    assert "join_code" in body
    assert len(body["join_code"]) == 8


# ---- Candidate filters (C3) -------------------------------------------------


async def test_filter_candidates_by_availability(
    client: AsyncClient,
    candidate_headers: dict[str, str],
    recruiter_headers: dict[str, str],
) -> None:
    org_r = await client.post(
        "/organizations", json={"name": "Filter Org"}, headers=recruiter_headers
    )
    org_id = org_r.json()["id"]
    await client.put(
        "/recruiters/me/profile", json={"organization_id": org_id}, headers=recruiter_headers
    )

    await client.put(
        "/candidates/me/profile",
        headers=candidate_headers,
        json={"availability_status": "available_now"},
    )

    inv = await client.post(
        f"/organizations/{org_id}/invitations",
        json={"candidate_email": "candidate@test.com"},
        headers=recruiter_headers,
    )
    token = inv.json()["token"]
    await client.post(f"/invitations/{token}/accept", headers=candidate_headers)

    r = await client.get(
        f"/organizations/{org_id}/candidates?availability_status=available_now",
        headers=recruiter_headers,
    )
    assert r.status_code == 200
    assert len(r.json()) == 1

    r2 = await client.get(
        f"/organizations/{org_id}/candidates?availability_status=not_available",
        headers=recruiter_headers,
    )
    assert r2.status_code == 200
    assert len(r2.json()) == 0


async def test_filter_candidates_by_skill(
    client: AsyncClient,
    candidate_headers: dict[str, str],
    recruiter_headers: dict[str, str],
) -> None:
    org_r = await client.post(
        "/organizations", json={"name": "Skill Org"}, headers=recruiter_headers
    )
    org_id = org_r.json()["id"]
    await client.put(
        "/recruiters/me/profile", json={"organization_id": org_id}, headers=recruiter_headers
    )

    ref_r = await client.post(
        "/skill-references",
        headers=candidate_headers,
        json={"name": "Python", "kind": "technical"},
    )
    ref_id = ref_r.json()["id"]
    await client.post(
        "/candidates/me/skills",
        headers=candidate_headers,
        json={"skill_ref_id": ref_id},
    )

    inv = await client.post(
        f"/organizations/{org_id}/invitations",
        json={"candidate_email": "candidate@test.com"},
        headers=recruiter_headers,
    )
    token = inv.json()["token"]
    await client.post(f"/invitations/{token}/accept", headers=candidate_headers)

    r = await client.get(
        f"/organizations/{org_id}/candidates?skill=python",
        headers=recruiter_headers,
    )
    assert r.status_code == 200
    assert len(r.json()) == 1

    r2 = await client.get(
        f"/organizations/{org_id}/candidates?skill=java",
        headers=recruiter_headers,
    )
    assert r2.status_code == 200
    assert len(r2.json()) == 0


# ---- Organization invitations -----------------------------------------------


async def test_list_org_invitations_returns_created_invitations(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    # Create org and link recruiter
    org_r = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "Inv Corp"}
    )
    org_id = org_r.json()["id"]
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )

    # Send an invitation
    await client.post(
        f"/organizations/{org_id}/invitations",
        headers=recruiter_headers,
        json={"candidate_email": "cand@test.com"},
    )

    r = await client.get(f"/organizations/{org_id}/invitations", headers=recruiter_headers)
    assert r.status_code == 200
    data = r.json()
    assert len(data) == 1
    assert data[0]["candidate_email"] == "cand@test.com"
    assert data[0]["status"] == "pending"


async def test_list_org_invitations_requires_membership(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    # Create a second recruiter that owns the org
    await client.post(
        "/auth/register",
        json={"email": "owner3@test.com", "password": "pass1234", "role": "recruiter"},
    )
    owner_login = await client.post(
        "/auth/login", json={"email": "owner3@test.com", "password": "pass1234"}
    )
    owner_headers = {"Authorization": f"Bearer {owner_login.json()['access_token']}"}

    org_r = await client.post("/organizations", headers=owner_headers, json={"name": "Other Corp"})
    other_org_id = org_r.json()["id"]

    # Original recruiter (not linked to this org) should get 403
    r = await client.get(f"/organizations/{other_org_id}/invitations", headers=recruiter_headers)
    assert r.status_code == 403


# ---- Atomic org creation + join/regen/members (Task 2) ---------------------


async def test_create_org_links_creator(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    """Creating an org should automatically set creator's organization_id."""
    r = await client.post(
        "/organizations",
        headers=recruiter_headers,
        json={"name": "Atomic Corp"},
    )
    assert r.status_code == 201
    org_id = r.json()["id"]

    profile_r = await client.get("/recruiters/me/profile", headers=recruiter_headers)
    assert profile_r.json()["organization_id"] == org_id


async def test_join_organization_by_code(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    second_recruiter_headers: dict[str, str],
) -> None:
    """A second recruiter can join an org via its join_code."""
    r = await client.post(
        "/organizations",
        headers=recruiter_headers,
        json={"name": "Join Test Corp"},
    )
    join_code = r.json()["join_code"]

    join_r = await client.post(
        "/organizations/join",
        headers=second_recruiter_headers,
        json={"code": join_code},
    )
    assert join_r.status_code == 200
    assert join_r.json()["organization_id"] == r.json()["id"]


async def test_join_invalid_code_returns_404(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    r = await client.post(
        "/organizations/join",
        headers=recruiter_headers,
        json={"code": "INVALID_CODE_XYZ"},
    )
    assert r.status_code == 404


async def test_regenerate_join_code(client: AsyncClient, recruiter_headers: dict[str, str]) -> None:
    r = await client.post(
        "/organizations",
        headers=recruiter_headers,
        json={"name": "Regen Corp"},
    )
    org_id = r.json()["id"]
    old_code = r.json()["join_code"]

    regen_r = await client.post(
        f"/organizations/{org_id}/regenerate-join-code",
        headers=recruiter_headers,
    )
    assert regen_r.status_code == 200
    assert regen_r.json()["join_code"] != old_code


async def test_list_members(client: AsyncClient, recruiter_headers: dict[str, str]) -> None:
    r = await client.post(
        "/organizations",
        headers=recruiter_headers,
        json={"name": "Members Corp"},
    )
    org_id = r.json()["id"]

    members_r = await client.get(
        f"/organizations/{org_id}/members",
        headers=recruiter_headers,
    )
    assert members_r.status_code == 200
    assert len(members_r.json()) >= 1
    assert "email" in members_r.json()[0]


async def test_filter_candidates_by_max_daily_rate(
    client: AsyncClient,
    candidate_headers: dict[str, str],
    recruiter_headers: dict[str, str],
) -> None:
    org_r = await client.post(
        "/organizations", json={"name": "Rate Org"}, headers=recruiter_headers
    )
    org_id = org_r.json()["id"]
    await client.put(
        "/recruiters/me/profile", json={"organization_id": org_id}, headers=recruiter_headers
    )

    await client.put("/candidates/me/profile", headers=candidate_headers, json={"daily_rate": 700})

    inv = await client.post(
        f"/organizations/{org_id}/invitations",
        json={"candidate_email": "candidate@test.com"},
        headers=recruiter_headers,
    )
    token = inv.json()["token"]
    await client.post(f"/invitations/{token}/accept", headers=candidate_headers)

    r = await client.get(
        f"/organizations/{org_id}/candidates?max_daily_rate=800",
        headers=recruiter_headers,
    )
    assert r.status_code == 200
    assert len(r.json()) == 1

    r2 = await client.get(
        f"/organizations/{org_id}/candidates?max_daily_rate=600",
        headers=recruiter_headers,
    )
    assert r2.status_code == 200
    assert len(r2.json()) == 0


async def test_filter_candidates_contract_type_includes_both(
    client: AsyncClient,
    candidate_headers: dict[str, str],
    recruiter_headers: dict[str, str],
) -> None:
    org_r = await client.post(
        "/organizations", json={"name": "Contract Org"}, headers=recruiter_headers
    )
    org_id = org_r.json()["id"]
    await client.put(
        "/recruiters/me/profile", json={"organization_id": org_id}, headers=recruiter_headers
    )

    # Le candidat est ouvert aux deux types de contrat.
    await client.put(
        "/candidates/me/profile",
        headers=candidate_headers,
        json={"contract_type": "both"},
    )

    inv = await client.post(
        f"/organizations/{org_id}/invitations",
        json={"candidate_email": "candidate@test.com"},
        headers=recruiter_headers,
    )
    token = inv.json()["token"]
    await client.post(f"/invitations/{token}/accept", headers=candidate_headers)

    for contract_filter in ("freelance", "cdi", "both"):
        r = await client.get(
            f"/organizations/{org_id}/candidates?contract_type={contract_filter}",
            headers=recruiter_headers,
        )
        assert r.status_code == 200
        assert len(r.json()) == 1, f"filter {contract_filter} should match a 'both' candidate"
