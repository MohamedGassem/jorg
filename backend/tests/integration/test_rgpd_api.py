# backend/tests/integration/test_rgpd_api.py
from datetime import UTC, datetime, timedelta

from httpx import AsyncClient
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from models.candidate_profile import CandidateProfile, Experience
from models.generated_document import GeneratedDocument
from models.invitation import AccessGrant, AccessGrantStatus, Invitation, InvitationStatus
from models.recruiter import Organization
from models.user import User


async def test_export_requires_auth(client: AsyncClient) -> None:
    r = await client.get("/candidates/me/export")
    assert r.status_code == 401


async def test_recruiter_cannot_export_candidate_data(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    r = await client.get("/candidates/me/export", headers=recruiter_headers)
    assert r.status_code == 403


async def test_export_empty_profile_returns_shell(
    client: AsyncClient, candidate_headers: dict[str, str]
) -> None:
    r = await client.get("/candidates/me/export", headers=candidate_headers)
    assert r.status_code == 200
    data = r.json()
    assert data["email"] == "candidate@test.com"
    assert data["role"] == "candidate"
    assert data["profile"] is None
    assert data["experiences"] == []
    assert data["skills"] == []
    assert data["education"] == []
    assert data["certifications"] == []
    assert data["languages"] == []
    assert data["access_grants"] == []
    assert data["generated_documents"] == []
    assert "exported_at" in data


async def test_delete_requires_auth(client: AsyncClient) -> None:
    r = await client.delete("/candidates/me")
    assert r.status_code == 401


async def test_recruiter_cannot_delete_candidate_account(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    r = await client.delete("/candidates/me", headers=recruiter_headers)
    assert r.status_code == 403


async def test_delete_removes_user_and_profile_cascade(
    client: AsyncClient,
    candidate_headers: dict[str, str],
    db_session: AsyncSession,
) -> None:
    # Donne du contenu à supprimer
    await client.put(
        "/candidates/me/profile",
        headers=candidate_headers,
        json={"first_name": "Alice", "last_name": "Dupont"},
    )
    await client.post(
        "/candidates/me/experiences",
        headers=candidate_headers,
        json={
            "client_name": "Acme",
            "role": "Dev",
            "start_date": "2024-01-01",
            "is_current": True,
        },
    )

    # Capture IDs before deletion so assertions are scoped to this candidate.
    user_q = await db_session.execute(select(User).where(User.email == "candidate@test.com"))
    candidate_user = user_q.scalar_one()
    candidate_user_id = candidate_user.id

    profile_q = await db_session.execute(
        select(CandidateProfile).where(CandidateProfile.user_id == candidate_user_id)
    )
    profile_before = profile_q.scalar_one()
    profile_id = profile_before.id

    r = await client.delete("/candidates/me", headers=candidate_headers)
    assert r.status_code == 204

    db_session.expire_all()

    user_check = await db_session.execute(select(User).where(User.email == "candidate@test.com"))
    assert user_check.scalar_one_or_none() is None

    profile_check = await db_session.execute(
        select(CandidateProfile).where(CandidateProfile.user_id == candidate_user_id)
    )
    assert profile_check.scalar_one_or_none() is None

    exp_check = await db_session.execute(
        select(Experience).where(Experience.profile_id == profile_id)
    )
    assert exp_check.scalar_one_or_none() is None


async def test_delete_anonymizes_access_grants_and_preserves_documents(
    client: AsyncClient,
    candidate_headers: dict[str, str],
    db_session: AsyncSession,
) -> None:
    # Construit manuellement un grant + document "existant"
    org = Organization(name="Conseil Co", slug="conseil-co")
    db_session.add(org)
    await db_session.flush()

    user_q = await db_session.execute(select(User).where(User.email == "candidate@test.com"))
    candidate_user = user_q.scalar_one()

    grant = AccessGrant(
        candidate_id=candidate_user.id,
        organization_id=org.id,
        status=AccessGrantStatus.ACTIVE,
        granted_at=datetime.now(UTC),
    )
    db_session.add(grant)
    await db_session.flush()

    doc = GeneratedDocument(
        access_grant_id=grant.id,
        template_id=None,
        generated_by_user_id=None,
        file_path="generated/old-doc.docx",
        file_format="docx",
    )
    db_session.add(doc)
    await db_session.commit()

    grant_id = grant.id
    doc_id = doc.id

    r = await client.delete("/candidates/me", headers=candidate_headers)
    assert r.status_code == 204

    db_session.expire_all()
    grant_q = await db_session.execute(select(AccessGrant).where(AccessGrant.id == grant_id))
    refreshed = grant_q.scalar_one()
    assert refreshed.candidate_id is None
    assert refreshed.status == AccessGrantStatus.REVOKED
    assert refreshed.revoked_at is not None

    doc_q = await db_session.execute(
        select(GeneratedDocument).where(GeneratedDocument.id == doc_id)
    )
    refreshed_doc = doc_q.scalar_one()
    assert refreshed_doc.access_grant_id == grant_id
    assert refreshed_doc.file_path == "generated/old-doc.docx"


async def test_delete_expires_pending_invitations(
    client: AsyncClient,
    candidate_headers: dict[str, str],
    db_session: AsyncSession,
) -> None:
    # Une orga + un recruteur + une invitation pending par email
    org = Organization(name="Other Co", slug="other-co")
    db_session.add(org)
    await db_session.flush()

    recruiter = User(email="rec@test.com", role="recruiter", hashed_password="x")
    db_session.add(recruiter)
    await db_session.flush()

    inv = Invitation(
        recruiter_id=recruiter.id,
        organization_id=org.id,
        candidate_email="candidate@test.com",
        candidate_id=None,
        token="tok-pending",
        status=InvitationStatus.PENDING,
        expires_at=datetime.now(UTC) + timedelta(days=30),
    )
    db_session.add(inv)
    await db_session.commit()

    inv_id = inv.id

    r = await client.delete("/candidates/me", headers=candidate_headers)
    assert r.status_code == 204

    db_session.expire_all()
    inv_q = await db_session.execute(select(Invitation).where(Invitation.id == inv_id))
    refreshed_inv = inv_q.scalar_one()
    assert refreshed_inv.status == InvitationStatus.EXPIRED


async def test_delete_does_not_expire_accepted_or_rejected_invitations(
    client: AsyncClient,
    candidate_headers: dict[str, str],
    db_session: AsyncSession,
) -> None:
    org = Organization(name="Fixed Co", slug="fixed-co")
    db_session.add(org)
    await db_session.flush()

    user_q = await db_session.execute(select(User).where(User.email == "candidate@test.com"))
    candidate_user = user_q.scalar_one()

    recruiter = User(email="rec-fixed@test.com", role="recruiter", hashed_password="x")
    db_session.add(recruiter)
    await db_session.flush()

    accepted_inv = Invitation(
        recruiter_id=recruiter.id,
        organization_id=org.id,
        candidate_email="candidate@test.com",
        candidate_id=candidate_user.id,
        token="tok-accepted",
        status=InvitationStatus.ACCEPTED,
        expires_at=datetime.now(UTC) + timedelta(days=30),
    )
    rejected_inv = Invitation(
        recruiter_id=recruiter.id,
        organization_id=org.id,
        candidate_email="candidate@test.com",
        candidate_id=candidate_user.id,
        token="tok-rejected",
        status=InvitationStatus.REJECTED,
        expires_at=datetime.now(UTC) + timedelta(days=30),
    )
    db_session.add(accepted_inv)
    db_session.add(rejected_inv)
    await db_session.commit()

    accepted_id = accepted_inv.id
    rejected_id = rejected_inv.id

    r = await client.delete("/candidates/me", headers=candidate_headers)
    assert r.status_code == 204

    db_session.expire_all()

    accepted_q = await db_session.execute(select(Invitation).where(Invitation.id == accepted_id))
    assert accepted_q.scalar_one().status == InvitationStatus.ACCEPTED

    rejected_q = await db_session.execute(select(Invitation).where(Invitation.id == rejected_id))
    assert rejected_q.scalar_one().status == InvitationStatus.REJECTED


async def test_export_includes_nested_candidate_skills(
    client: AsyncClient, candidate_headers: dict[str, str]
) -> None:
    ref1_r = await client.post(
        "/skill-references",
        headers=candidate_headers,
        json={"name": "RGPDSkill1", "kind": "technical"},
    )
    ref1_id = ref1_r.json()["id"]

    ref2_r = await client.post(
        "/skill-references",
        headers=candidate_headers,
        json={"name": "RGPDSkill2", "kind": "tool"},
    )
    ref2_id = ref2_r.json()["id"]

    await client.post(
        "/candidates/me/skills",
        headers=candidate_headers,
        json={"skill_ref_id": ref1_id, "featured": True},
    )
    await client.post(
        "/candidates/me/skills",
        headers=candidate_headers,
        json={"skill_ref_id": ref2_id},
    )

    exp_r = await client.post(
        "/candidates/me/experiences",
        headers=candidate_headers,
        json={"client_name": "RGPDCorp", "role": "Dev", "start_date": "2022-01-01"},
    )
    exp_id = exp_r.json()["id"]
    await client.post(
        f"/candidates/me/experiences/{exp_id}/skill-usages",
        headers=candidate_headers,
        json={"skill_ref_id": ref1_id, "intensity": "primary"},
    )

    r = await client.get("/candidates/me/export", headers=candidate_headers)
    assert r.status_code == 200
    data = r.json()

    assert len(data["skills"]) >= 2
    skill_names = [s["skill_ref"]["name"] for s in data["skills"]]
    assert "RGPDSkill1" in skill_names
    assert "RGPDSkill2" in skill_names

    s1 = next(s for s in data["skills"] if s["skill_ref"]["name"] == "RGPDSkill1")
    assert "id" in s1["skill_ref"]
    assert "kind" in s1["skill_ref"]
    assert s1["featured"] is True


# ---- Recruiter RGPD ---------------------------------------------------------


async def test_recruiter_export_requires_auth(client: AsyncClient) -> None:
    r = await client.get("/recruiters/me/export")
    assert r.status_code == 401


async def test_candidate_cannot_export_recruiter_data(
    client: AsyncClient, candidate_headers: dict[str, str]
) -> None:
    r = await client.get("/recruiters/me/export", headers=candidate_headers)
    assert r.status_code == 403


async def test_recruiter_export_returns_profile_and_organization(
    client: AsyncClient, recruiter_headers: dict[str, str]
) -> None:
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"first_name": "Bob", "last_name": "Martin", "job_title": "Talent"},
    )
    org = await client.post("/organizations", headers=recruiter_headers, json={"name": "Export Co"})
    org_id = org.json()["id"]
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )

    r = await client.get("/recruiters/me/export", headers=recruiter_headers)
    assert r.status_code == 200
    data = r.json()
    assert data["email"] == "recruiter@test.com"
    assert data["role"] == "recruiter"
    assert data["first_name"] == "Bob"
    assert data["last_name"] == "Martin"
    assert data["job_title"] == "Talent"
    assert data["organization"]["id"] == org_id
    assert data["organization"]["name"] == "Export Co"
    assert data["generated_documents"] == []
    assert "exported_at" in data


async def test_recruiter_delete_requires_auth(client: AsyncClient) -> None:
    r = await client.delete("/recruiters/me")
    assert r.status_code == 401


async def test_candidate_cannot_delete_recruiter_account(
    client: AsyncClient, candidate_headers: dict[str, str]
) -> None:
    r = await client.delete("/recruiters/me", headers=candidate_headers)
    assert r.status_code == 403


async def test_recruiter_delete_removes_user_and_invalidates_session(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    db_session: AsyncSession,
) -> None:
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"first_name": "Gone", "last_name": "Soon"},
    )

    user_q = await db_session.execute(select(User).where(User.email == "recruiter@test.com"))
    recruiter_user = user_q.scalar_one()
    recruiter_user_id = recruiter_user.id

    r = await client.delete("/recruiters/me", headers=recruiter_headers)
    assert r.status_code == 204

    db_session.expire_all()

    user_check = await db_session.execute(select(User).where(User.email == "recruiter@test.com"))
    assert user_check.scalar_one_or_none() is None

    from models.recruiter import RecruiterProfile

    profile_check = await db_session.execute(
        select(RecruiterProfile).where(RecruiterProfile.user_id == recruiter_user_id)
    )
    assert profile_check.scalar_one_or_none() is None

    # A subsequent authenticated call with the old token now fails.
    after = await client.get("/recruiters/me/profile", headers=recruiter_headers)
    assert after.status_code == 401


async def test_recruiter_delete_preserves_org_and_nulls_generated_docs(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    db_session: AsyncSession,
) -> None:
    org_r = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "Survivor Co"}
    )
    org_id = org_r.json()["id"]

    user_q = await db_session.execute(select(User).where(User.email == "recruiter@test.com"))
    recruiter_user = user_q.scalar_one()

    doc = GeneratedDocument(
        access_grant_id=None,
        template_id=None,
        generated_by_user_id=recruiter_user.id,
        file_path="generated/rec-doc.docx",
        file_format="docx",
    )
    db_session.add(doc)
    await db_session.commit()
    doc_id = doc.id

    r = await client.delete("/recruiters/me", headers=recruiter_headers)
    assert r.status_code == 204

    db_session.expire_all()

    org_check = await db_session.execute(select(Organization).where(Organization.id == org_id))
    assert org_check.scalar_one_or_none() is not None

    doc_q = await db_session.execute(
        select(GeneratedDocument).where(GeneratedDocument.id == doc_id)
    )
    refreshed_doc = doc_q.scalar_one()
    assert refreshed_doc.generated_by_user_id is None
    assert refreshed_doc.file_path == "generated/rec-doc.docx"


async def test_recruiter_delete_preserves_owned_template(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    db_session: AsyncSession,
) -> None:
    """Deleting a recruiter who created a template must not be blocked by its FK.

    templates.created_by_user_id is ON DELETE SET NULL: the template survives for
    the organization, its creator is anonymised.
    """
    import io
    from uuid import UUID

    from docx import Document  # type: ignore[import-untyped,unused-ignore]

    from models.template import Template

    org_r = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "Template Owner Co"}
    )
    org_id = org_r.json()["id"]

    doc = Document()
    doc.add_paragraph("Nom: {{last_name}}")
    buf = io.BytesIO()
    doc.save(buf)
    up = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "Owned"},
        files={
            "file": (
                "t.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert up.status_code == 201, up.text
    template_id = UUID(up.json()["id"])

    r = await client.delete("/recruiters/me", headers=recruiter_headers)
    assert r.status_code == 204

    db_session.expire_all()
    t_q = await db_session.execute(select(Template).where(Template.id == template_id))
    template = t_q.scalar_one()
    assert template.created_by_user_id is None
