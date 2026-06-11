# backend/tests/integration/test_generation_api.py
import io

from docx import Document  # type: ignore[import-untyped,unused-ignore]
from httpx import AsyncClient

# ---- helpers ----------------------------------------------------------------


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _setup_org_with_grant(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> tuple[str, str]:
    """Create org, link recruiter, invite+accept candidate. Returns (org_id, candidate_id)."""
    org = await client.post("/organizations", headers=recruiter_headers, json={"name": "GenCorp"})
    org_id: str = org.json()["id"]
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )
    # Get candidate id from profile
    profile = await client.get("/candidates/me/profile", headers=candidate_headers)
    candidate_id: str = profile.json()["user_id"]

    inv = await client.post(
        f"/organizations/{org_id}/invitations",
        headers=recruiter_headers,
        json={"candidate_email": "candidate@test.com"},
    )
    token = inv.json()["token"]
    await client.post(f"/invitations/{token}/accept", headers=candidate_headers)
    return org_id, candidate_id


async def _upload_valid_template(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    org_id: str,
) -> str:
    """Upload a template with standard docxtpl placeholders. Returns template_id."""
    docx_bytes = _make_docx_bytes(["Nom: {{last_name}}", "Titre: {{title}}"])
    r = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "CV Template"},
        files={
            "file": (
                "t.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 201, r.text
    assert r.json()["is_valid"] is True
    return str(r.json()["id"])


# ---- generate ---------------------------------------------------------------


async def test_recruiter_generates_document(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)
    template_id = await _upload_valid_template(client, recruiter_headers, org_id)

    r = await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )
    assert r.status_code == 201
    data = r.json()
    assert data["file_format"] == "docx"
    assert "id" in data
    assert data["template_id"] == template_id


async def test_builtin_template_preview_downloads_mock_docx(
    client: AsyncClient,
    candidate_headers: dict[str, str],
) -> None:
    templates = await client.get("/templates/builtin", headers=candidate_headers)
    assert templates.status_code == 200
    assert {t["key"] for t in templates.json()} == {
        "compact_esn",
        "dossier_technique",
        "profil_premium",
    }

    preview = await client.get(
        "/templates/builtin/compact_esn/preview",
        headers=candidate_headers,
    )
    assert preview.status_code == 200
    doc = Document(io.BytesIO(preview.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "joris" in text.lower()
    assert "{{" not in text


async def test_candidate_generates_own_document_from_builtin_template(
    client: AsyncClient,
    candidate_headers: dict[str, str],
) -> None:
    await client.get("/candidates/me/profile", headers=candidate_headers)

    r = await client.post(
        "/candidates/me/generate",
        headers=candidate_headers,
        json={"system_template_key": "compact_esn", "format": "docx"},
    )
    assert r.status_code == 201, r.text
    data = r.json()
    assert data["access_grant_id"] is None
    assert data["template_id"] is None
    assert data["template_name"] == "Synthèse"

    download = await client.get(f"/documents/{data['id']}/download", headers=candidate_headers)
    assert download.status_code == 200
    assert download.content.startswith(b"PK")


async def test_cannot_generate_without_access_grant(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    # Setup org + template but NO invitation/grant
    org = await client.post(
        "/organizations", headers=recruiter_headers, json={"name": "NoGrant Corp"}
    )
    org_id: str = org.json()["id"]
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"organization_id": org_id},
    )
    template_id = await _upload_valid_template(client, recruiter_headers, org_id)
    profile = await client.get("/candidates/me/profile", headers=candidate_headers)
    candidate_id = profile.json()["user_id"]

    r = await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )
    assert r.status_code == 403


async def test_cannot_generate_with_invalid_template(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)
    # Unknown placeholders are unsupported and make the template invalid.
    docx_bytes = _make_docx_bytes(["{{NOM}} {{UNMAPPED}}"])
    r = await client.post(
        f"/organizations/{org_id}/templates",
        headers=recruiter_headers,
        data={"name": "Bad Template"},
        files={
            "file": (
                "t.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    template_id = r.json()["id"]
    assert r.json()["is_valid"] is False

    r2 = await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )
    assert r2.status_code == 422


async def test_candidate_cannot_trigger_generation(
    client: AsyncClient,
    candidate_headers: dict[str, str],
) -> None:
    r = await client.post(
        "/organizations/00000000-0000-0000-0000-000000000000/generate",
        headers=candidate_headers,
        json={
            "candidate_id": "00000000-0000-0000-0000-000000000001",
            "template_id": "00000000-0000-0000-0000-000000000002",
            "format": "docx",
        },
    )
    assert r.status_code == 403


async def test_non_member_recruiter_cannot_generate(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    second_recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    """A recruiter who is not a member of the org gets 403 on POST /organizations/{org_id}/generate."""
    # recruiter_headers creates and owns the org; second_recruiter_headers is NOT a member
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)
    template_id = await _upload_valid_template(client, recruiter_headers, org_id)

    r = await client.post(
        f"/organizations/{org_id}/generate",
        headers=second_recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )
    assert r.status_code == 403


# ---- history ----------------------------------------------------------------


async def test_candidate_history_lists_generated_docs(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)
    template_id = await _upload_valid_template(client, recruiter_headers, org_id)
    await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )

    r = await client.get("/candidates/me/documents", headers=candidate_headers)
    assert r.status_code == 200
    assert len(r.json()) >= 1


async def test_recruiter_org_history(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)
    template_id = await _upload_valid_template(client, recruiter_headers, org_id)
    await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )

    r = await client.get(f"/organizations/{org_id}/documents", headers=recruiter_headers)
    assert r.status_code == 200
    assert len(r.json()) >= 1


# ---- download ---------------------------------------------------------------


async def test_candidate_documents_no_file_path(
    client: AsyncClient,
    candidate_headers: dict[str, str],
) -> None:
    """Candidate document list must not expose file_path."""
    r = await client.get("/candidates/me/documents", headers=candidate_headers)
    assert r.status_code == 200
    docs = r.json()
    for doc in docs:
        assert "file_path" not in doc


async def test_candidate_documents_have_org_and_template_name(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    """Candidate document list includes organization_name and template_name."""
    # This test requires a real document to exist. If none exist, it's a no-op.
    r = await client.get("/candidates/me/documents", headers=candidate_headers)
    assert r.status_code == 200
    docs = r.json()
    for doc in docs:
        assert "organization_name" in doc
        assert "template_name" in doc
        assert "file_path" not in doc


# ---- recruiter name in candidate view ---------------------------------------


async def test_candidate_documents_include_recruiter_name(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    """Candidate document view includes recruiter first_name and last_name."""
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)

    # Set recruiter name on their profile
    await client.put(
        "/recruiters/me/profile",
        headers=recruiter_headers,
        json={"first_name": "Alice", "last_name": "Smith"},
    )

    template_id = await _upload_valid_template(client, recruiter_headers, org_id)
    await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )

    r = await client.get("/candidates/me/documents", headers=candidate_headers)
    assert r.status_code == 200
    docs = r.json()
    assert len(docs) >= 1
    doc = docs[0]
    assert doc["recruiter_first_name"] == "Alice"
    assert doc["recruiter_last_name"] == "Smith"


# ---- recruiter org document view --------------------------------------------


async def test_org_documents_include_candidate_name(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    """Recruiter org document list includes candidate_first_name and candidate_last_name."""
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)

    # Set candidate name on their profile
    await client.put(
        "/candidates/me/profile",
        headers=candidate_headers,
        json={"first_name": "Bob", "last_name": "Jones"},
    )

    template_id = await _upload_valid_template(client, recruiter_headers, org_id)
    await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )

    r = await client.get(f"/organizations/{org_id}/documents", headers=recruiter_headers)
    assert r.status_code == 200
    docs = r.json()
    assert len(docs) >= 1
    doc = docs[0]
    assert doc["candidate_first_name"] == "Bob"
    assert doc["candidate_last_name"] == "Jones"


# ---- download ---------------------------------------------------------------


async def test_org_documents_no_duplicate_when_candidate_on_multiple_shortlists(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    """Candidate shortlisted for multiple opportunities must not duplicate generated documents."""
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)
    template_id = await _upload_valid_template(client, recruiter_headers, org_id)

    # Generate one document
    await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )

    # Create two opportunities and add the candidate to both shortlists
    for title in ("Opportunity Alpha", "Opportunity Beta"):
        opp_r = await client.post(
            f"/organizations/{org_id}/opportunities",
            headers=recruiter_headers,
            json={"title": title},
        )
        assert opp_r.status_code == 201
        opp_id = opp_r.json()["id"]
        sl_r = await client.post(
            f"/organizations/{org_id}/opportunities/{opp_id}/candidates",
            headers=recruiter_headers,
            json={"candidate_id": candidate_id},
        )
        assert sl_r.status_code == 201

    r = await client.get(f"/organizations/{org_id}/documents", headers=recruiter_headers)
    assert r.status_code == 200
    docs = r.json()
    # Despite being on 2 shortlists, only 1 document should appear
    assert len(docs) == 1


async def test_download_generated_document(
    client: AsyncClient,
    recruiter_headers: dict[str, str],
    candidate_headers: dict[str, str],
) -> None:
    org_id, candidate_id = await _setup_org_with_grant(client, recruiter_headers, candidate_headers)
    template_id = await _upload_valid_template(client, recruiter_headers, org_id)
    gen = await client.post(
        f"/organizations/{org_id}/generate",
        headers=recruiter_headers,
        json={"candidate_id": candidate_id, "template_id": template_id, "format": "docx"},
    )
    doc_id = gen.json()["id"]

    r = await client.get(f"/documents/{doc_id}/download", headers=recruiter_headers)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument")
