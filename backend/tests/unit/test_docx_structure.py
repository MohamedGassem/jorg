# backend/tests/unit/test_docx_structure.py
from pathlib import Path

from docx import Document

from services.documents.docx_structure import extract_structure


def _sample(tmp_path: Path) -> str:
    doc = Document()
    doc.add_paragraph("Jean Dupont")
    doc.add_paragraph("Consultant senior")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Client"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "ACME"
    table.cell(1, 1).text = "Dev"
    path = tmp_path / "s.docx"
    doc.save(str(path))
    return str(path)


def test_extract_structure_indexes_paragraphs_and_tables(tmp_path: Path) -> None:
    structure = extract_structure(_sample(tmp_path))
    assert structure["paragraphs"][0] == {"index": 0, "text": "Jean Dupont"}
    assert structure["paragraphs"][1] == {"index": 1, "text": "Consultant senior"}
    table = structure["tables"][0]
    assert table["index"] == 0
    assert table["rows"][0] == {"index": 0, "cells": ["Client", "Role"]}
    assert table["rows"][1] == {"index": 1, "cells": ["ACME", "Dev"]}
