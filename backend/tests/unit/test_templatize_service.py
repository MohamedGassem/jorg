# backend/tests/unit/test_templatize_service.py
from pathlib import Path
from typing import Any

from docx import Document

from services.documents.templatize_ops import ReplaceTextOp, TemplatizePlan
from services.documents.templatize_service import TemplatizeOutcome, run_templatize_pipeline


def _example_docx(tmp_path: Path) -> str:
    doc = Document()
    doc.add_paragraph("Jean Dupont")
    doc.add_paragraph("Consultant senior")
    path = tmp_path / "model.docx"
    doc.save(str(path))
    return str(path)


async def test_pipeline_applies_plan_and_renders(tmp_path: Path, monkeypatch: Any) -> None:
    plan = TemplatizePlan(
        operations=[
            ReplaceTextOp(
                op="replace_text",
                target={"kind": "paragraph", "paragraph": 0},
                find="Jean Dupont",
                placeholder="{{first_name}} {{last_name}}",
            )
        ]
    )

    async def fake_request_plan(client: Any, **kwargs: Any) -> TemplatizePlan:
        return plan

    monkeypatch.setattr("services.documents.templatize_service.request_plan", fake_request_plan)
    outcome = await run_templatize_pipeline(object(), "claude-opus-4-8", _example_docx(tmp_path))
    assert isinstance(outcome, TemplatizeOutcome)
    assert outcome.render_error is None
    assert outcome.report["mappings"] == [
        {"find": "Jean Dupont", "placeholder": "{{first_name}} {{last_name}}"}
    ]
    assert outcome.docx_bytes[:2] == b"PK"


async def test_pipeline_retries_once_on_render_failure(tmp_path: Path, monkeypatch: Any) -> None:
    bad_plan = TemplatizePlan(
        operations=[
            ReplaceTextOp(
                op="replace_text",
                target={"kind": "paragraph", "paragraph": 0},
                find="Jean Dupont",
                placeholder="{%p endfor %}",  # casse le rendu Jinja
            )
        ]
    )
    calls: list[str | None] = []

    async def fake_request_plan(
        client: Any, *, render_errors: str | None, **kwargs: Any
    ) -> TemplatizePlan:
        calls.append(render_errors)
        return bad_plan

    monkeypatch.setattr("services.documents.templatize_service.request_plan", fake_request_plan)
    outcome = await run_templatize_pipeline(object(), "claude-opus-4-8", _example_docx(tmp_path))
    assert len(calls) == 2  # tentative initiale + 1 retry
    assert calls[0] is None and calls[1] is not None
    assert outcome.render_error is not None  # brouillon conserve avec rapport d'echec
