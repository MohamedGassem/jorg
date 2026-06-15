# backend/tests/unit/test_templatize_ops.py
import io
from collections.abc import Callable
from pathlib import Path

from docx import Document

from services.documents.templatize_ops import (
    DeleteBlockOp,
    FlagResidualOp,
    ReplaceTextOp,
    TemplatizePlan,
    WrapParagraphsLoopOp,
    WrapTableRowsLoopOp,
    apply_operations,
)


def _doc_path(tmp_path: Path, build: Callable[[object], None]) -> str:
    doc = Document()
    build(doc)
    path = tmp_path / "t.docx"
    doc.save(str(path))
    return str(path)


def _load(result_bytes: bytes) -> object:
    return Document(io.BytesIO(result_bytes))


def test_replace_text_in_paragraph_preserves_other_runs(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        p = doc.add_paragraph()  # type: ignore[attr-defined]
        p.add_run("Nom : ").bold = True
        p.add_run("Jean Dupont")

    path = _doc_path(tmp_path, build)
    plan = TemplatizePlan(
        operations=[
            ReplaceTextOp(
                op="replace_text",
                target={"kind": "paragraph", "paragraph": 0},
                find="Jean Dupont",
                placeholder="{{first_name}} {{last_name}}",
            )
        ]
    )
    result = apply_operations(path, plan)
    doc = _load(result.docx_bytes)
    assert doc.paragraphs[0].text == "Nom : {{first_name}} {{last_name}}"  # type: ignore[attr-defined]
    assert doc.paragraphs[0].runs[0].bold is True  # type: ignore[attr-defined]
    assert result.rejected == []


def test_replace_text_in_table_cell(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        table = doc.add_table(rows=1, cols=2)  # type: ignore[attr-defined]
        table.cell(0, 0).text = "ACME Corp"
        table.cell(0, 1).text = "Lead Dev"

    path = _doc_path(tmp_path, build)
    plan = TemplatizePlan(
        operations=[
            ReplaceTextOp(
                op="replace_text",
                target={"kind": "cell", "table": 0, "row": 0, "cell": 1},
                find="Lead Dev",
                placeholder="{{exp.role}}",
            )
        ]
    )
    result = apply_operations(path, plan)
    doc = _load(result.docx_bytes)
    assert doc.tables[0].cell(0, 1).text == "{{exp.role}}"  # type: ignore[attr-defined]


def test_wrap_paragraphs_loop_inserts_tags_around_block(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        doc.add_paragraph("Titre")  # type: ignore[attr-defined]
        doc.add_paragraph("ACME - Lead Dev")  # type: ignore[attr-defined]
        doc.add_paragraph("Description mission")  # type: ignore[attr-defined]

    path = _doc_path(tmp_path, build)
    plan = TemplatizePlan(
        operations=[
            WrapParagraphsLoopOp(
                op="wrap_paragraphs_loop",
                start_paragraph=1,
                end_paragraph=2,
                loop_var="exp",
                collection="experiences",
            )
        ]
    )
    result = apply_operations(path, plan)
    texts = [p.text for p in _load(result.docx_bytes).paragraphs]  # type: ignore[attr-defined]
    assert texts == [
        "Titre",
        "{%p for exp in experiences %}",
        "ACME - Lead Dev",
        "Description mission",
        "{%p endfor %}",
    ]


def test_wrap_table_rows_loop_inserts_dedicated_rows(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        table = doc.add_table(rows=2, cols=2)  # type: ignore[attr-defined]
        table.cell(0, 0).text = "Client"
        table.cell(1, 0).text = "ACME"

    path = _doc_path(tmp_path, build)
    plan = TemplatizePlan(
        operations=[
            WrapTableRowsLoopOp(
                op="wrap_table_rows_loop",
                table=0,
                start_row=1,
                end_row=1,
                loop_var="exp",
                collection="experiences",
            )
        ]
    )
    result = apply_operations(path, plan)
    table = _load(result.docx_bytes).tables[0]  # type: ignore[attr-defined]
    assert len(table.rows) == 4
    assert table.cell(1, 0).text == "{%tr for exp in experiences %}"
    assert table.cell(2, 0).text == "ACME"
    assert table.cell(3, 0).text == "{%tr endfor %}"


def test_delete_block_removes_paragraphs_and_rows(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        doc.add_paragraph("Garde")  # type: ignore[attr-defined]
        doc.add_paragraph("Exemple 2 a supprimer")  # type: ignore[attr-defined]
        table = doc.add_table(rows=3, cols=1)  # type: ignore[attr-defined]
        table.cell(0, 0).text = "Entete"
        table.cell(1, 0).text = "Gabarit"
        table.cell(2, 0).text = "Doublon"

    path = _doc_path(tmp_path, build)
    plan = TemplatizePlan(
        operations=[
            DeleteBlockOp(
                op="delete_block",
                target={"kind": "paragraphs", "start_paragraph": 1, "end_paragraph": 1},
            ),
            DeleteBlockOp(
                op="delete_block",
                target={"kind": "table_rows", "table": 0, "start_row": 2, "end_row": 2},
            ),
        ]
    )
    result = apply_operations(path, plan)
    doc = _load(result.docx_bytes)
    assert [p.text for p in doc.paragraphs if p.text] == ["Garde"]  # type: ignore[attr-defined]
    assert len(doc.tables[0].rows) == 2  # type: ignore[attr-defined]


def test_invalid_targets_are_rejected_not_raised(tmp_path: Path) -> None:
    def build(doc: object) -> None:
        doc.add_paragraph("Seul paragraphe")  # type: ignore[attr-defined]

    path = _doc_path(tmp_path, build)
    plan = TemplatizePlan(
        operations=[
            ReplaceTextOp(
                op="replace_text",
                target={"kind": "paragraph", "paragraph": 99},
                find="x",
                placeholder="{{x}}",
            ),
            ReplaceTextOp(
                op="replace_text",
                target={"kind": "paragraph", "paragraph": 0},
                find="introuvable",
                placeholder="{{x}}",
            ),
            FlagResidualOp(
                op="flag_residual",
                target={"kind": "paragraph", "paragraph": 0},
                note="Accroche fictive a relire",
            ),
        ]
    )
    result = apply_operations(path, plan)
    assert len(result.rejected) == 2
    assert result.residual_flags == ["Accroche fictive a relire"]
    assert _load(result.docx_bytes).paragraphs[0].text == "Seul paragraphe"  # type: ignore[attr-defined]
