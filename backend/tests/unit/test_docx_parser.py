# backend/tests/unit/test_docx_parser.py
import tempfile

from docx import Document  # type: ignore[import-untyped,unused-ignore]
from docx.shared import Inches

from services.documents.docx_parser import extract_placeholders


def _make_docx(paragraphs: list[str]) -> str:
    """Create a temporary .docx file with the given paragraphs, return path."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        return tmp.name


def test_extract_simple_placeholders() -> None:
    path = _make_docx(["Nom: {{NOM}}", "Prénom: {{PRENOM}}", "Titre: {{TITRE}}"])
    result = extract_placeholders(path)
    assert "{{NOM}}" in result
    assert "{{PRENOM}}" in result
    assert "{{TITRE}}" in result
    assert len(result) == 3


def test_extract_deduplicates_repeated_placeholders() -> None:
    path = _make_docx(["{{NOM}} et {{NOM}} encore {{NOM}}"])
    result = extract_placeholders(path)
    assert result.count("{{NOM}}") == 1


def test_extract_block_markers_are_excluded() -> None:
    """Block markers like {{#EXPERIENCES}} / {{/EXPERIENCES}} are mustache
    control syntax, not data placeholders — they must not be returned."""
    path = _make_docx(["{{#EXPERIENCES}}", "{{EXP_CLIENT}}", "{{/EXPERIENCES}}"])
    result = extract_placeholders(path)
    assert "{{#EXPERIENCES}}" not in result
    assert "{{/EXPERIENCES}}" not in result
    assert "{{EXP_CLIENT}}" in result
    assert result == ["{{EXP_CLIENT}}"]


def test_extract_block_markers_excluded_for_any_block_name() -> None:
    """The filter must apply to any {{#NAME}} / {{/NAME}} pair, not just EXPERIENCES."""
    path = _make_docx(["{{#SKILLS}}", "{{SKILL_NAME}}", "{{/SKILLS}}"])
    result = extract_placeholders(path)
    assert result == ["{{SKILL_NAME}}"]


def test_extract_empty_document_returns_empty_list() -> None:
    path = _make_docx(["No placeholders here."])
    result = extract_placeholders(path)
    assert result == []


def test_extract_preserves_first_occurrence_order() -> None:
    path = _make_docx(["{{A}} {{B}} {{C}} {{A}}"])
    result = extract_placeholders(path)
    assert result == ["{{A}}", "{{B}}", "{{C}}"]


# ---- Headers and footers -------------------------------------------------------


def test_placeholder_in_header_is_detected() -> None:
    doc = Document()
    doc.add_paragraph("body only")
    doc.sections[0].header.paragraphs[0].text = "{{first_name}}"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name
    result = extract_placeholders(path)
    assert "{{first_name}}" in result


def test_placeholder_in_footer_is_detected() -> None:
    doc = Document()
    doc.add_paragraph("body only")
    doc.sections[0].footer.paragraphs[0].text = "{{last_name}}"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name
    result = extract_placeholders(path)
    assert "{{last_name}}" in result


def test_header_placeholder_deduplicated_with_body() -> None:
    """Same placeholder in header and body counts only once."""
    doc = Document()
    doc.add_paragraph("{{title}}")
    doc.sections[0].header.paragraphs[0].text = "{{title}}"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name
    result = extract_placeholders(path)
    assert result.count("{{title}}") == 1


# ---- docxtpl / Jinja2 syntax (new templates) ---------------------------------


def test_jinja2_block_tags_not_matched() -> None:
    """Jinja2 block tags {%p for ... %} use {%, not {{, so they are never
    matched by the placeholder regex and never appear in detected placeholders."""
    path = _make_docx(
        [
            "{%p for exp in experiences %}",
            "{{exp.role}}",
            "{%p endfor %}",
        ]
    )
    result = extract_placeholders(path)
    assert not any("{%p" in r for r in result)
    assert not any("{%tr" in r for r in result)


def test_jinja2_loop_variables_are_excluded() -> None:
    """Loop variables {{exp.*}} and {{sk.*}} are not standalone mappable fields
    — they are resolved inside {%p for ... %} blocks."""
    path = _make_docx(
        [
            "{%p for exp in experiences %}",
            "{{exp.client_name}} — {{exp.role}}",
            "{{exp.start_date}} / {{exp.technologies}}",
            "{%p endfor %}",
            "{%p for sk in skills %}",
            "{{sk.name}} ({{sk.category}})",
            "{%p endfor %}",
        ]
    )
    result = extract_placeholders(path)
    assert not any(r.startswith("{{exp.") for r in result)
    assert not any(r.startswith("{{sk.") for r in result)


def test_profile_fields_detected_alongside_jinja2_blocks() -> None:
    """Top-level profile fields are still detected even when the template also
    uses Jinja2 blocks for experiences or skills."""
    path = _make_docx(
        [
            "{{last_name}} {{first_name}}",
            "{{title}}",
            "{%p for exp in experiences %}",
            "{{exp.role}} chez {{exp.client_name}}",
            "{%p endfor %}",
        ]
    )
    result = extract_placeholders(path)
    assert "{{last_name}}" in result
    assert "{{first_name}}" in result
    assert "{{title}}" in result
    assert not any(r.startswith("{{exp.") for r in result)


def test_standard_field_names_detected() -> None:
    """All standard profile field names used directly in a docxtpl template
    are surfaced as detected placeholders."""
    path = _make_docx(
        [
            "{{first_name}} {{last_name}}",
            "{{title}} — {{location}}",
            "TJM : {{daily_rate}} € | Salaire : {{annual_salary}} €",
            "{{availability_status}} / {{work_mode}}",
        ]
    )
    result = extract_placeholders(path)
    for field in [
        "{{first_name}}",
        "{{last_name}}",
        "{{title}}",
        "{{location}}",
        "{{daily_rate}}",
        "{{annual_salary}}",
        "{{availability_status}}",
        "{{work_mode}}",
    ]:
        assert field in result, f"{field} not detected"


# ---- Table cells inside headers/footers ------------------------------------


def test_placeholder_in_header_table_cell_is_detected() -> None:
    """A placeholder inside a table cell in the header must be detected."""
    doc = Document()
    doc.add_paragraph("body")
    header = doc.sections[0].header
    table = header.add_table(rows=1, cols=2, width=Inches(6))
    table.rows[0].cells[0].paragraphs[0].text = "{{first_name}}"
    table.rows[0].cells[1].paragraphs[0].text = "{{last_name}}"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name
    result = extract_placeholders(path)
    assert "{{first_name}}" in result
    assert "{{last_name}}" in result


def test_placeholder_in_footer_table_cell_is_detected() -> None:
    """A placeholder inside a table cell in the footer must be detected."""
    doc = Document()
    doc.add_paragraph("body")
    footer = doc.sections[0].footer
    table = footer.add_table(rows=1, cols=1, width=Inches(6))
    table.rows[0].cells[0].paragraphs[0].text = "{{title}}"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name
    result = extract_placeholders(path)
    assert "{{title}}" in result


# ---- even_page header/footer -----------------------------------------------


def test_placeholder_in_even_page_header_is_detected() -> None:
    """Placeholders in the even-page header are detected."""
    doc = Document()
    doc.add_paragraph("body")
    doc.sections[0].different_first_page_header_footer = True
    even_hdr = doc.sections[0].even_page_header
    even_hdr.paragraphs[0].text = "{{location}}"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name
    result = extract_placeholders(path)
    assert "{{location}}" in result


def test_placeholder_in_even_page_footer_is_detected() -> None:
    """Placeholders in the even-page footer are detected."""
    doc = Document()
    doc.add_paragraph("body")
    even_ftr = doc.sections[0].even_page_footer
    even_ftr.paragraphs[0].text = "{{phone}}"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name
    result = extract_placeholders(path)
    assert "{{phone}}" in result


# ---- Custom loop variable names NOT filtered --------------------------------


def test_custom_loop_variable_name_is_not_filtered() -> None:
    """The _LOOP_VAR_RE filter only applies to exp.* and sk.* — the documented
    loop variable names. A placeholder like {{item.foo}} does NOT match the
    filter and will appear in detected_placeholders.

    This is the expected behaviour: recruiters using non-standard loop var
    names will see those placeholders in detected_placeholders and be alerted
    that they are unmapped (and must use exp/sk instead).
    """
    path = _make_docx(["{{item.client_name}}", "{{entry.role}}"])
    result = extract_placeholders(path)
    assert "{{item.client_name}}" in result
    assert "{{entry.role}}" in result


# ---- Table cell placeholder detection in document body ----------------------


def test_placeholder_in_table_cell_body_is_detected() -> None:
    """Placeholders inside table cells in the document body are detected."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].text = "{{first_name}}"
    table.rows[0].cells[1].paragraphs[0].text = "{{last_name}}"
    table.rows[1].cells[0].paragraphs[0].text = "{{title}}"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name
    result = extract_placeholders(path)
    assert "{{first_name}}" in result
    assert "{{last_name}}" in result
    assert "{{title}}" in result
