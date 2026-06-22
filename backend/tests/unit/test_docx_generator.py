# backend/tests/unit/test_docx_generator.py
"""Integration tests for generate_document — new docxtpl syntax."""

import io
import tempfile
from datetime import date
from unittest.mock import MagicMock

from docx import Document

from models.candidate_profile import LanguageLevel
from models.skill import SkillKind, UsageIntensity
from services.documents.docx_engine import generate_document


def _make_docx_path(paragraphs: list[str]) -> str:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        return tmp.name


def _mock_profile(**kwargs: object) -> MagicMock:
    defaults: dict[str, object] = {
        "first_name": "Alice",
        "last_name": "Martin",
        "title": "Software Engineer",
        "summary": "Senior developer with 8 years of experience",
        "phone": "0601020304",
        "email_contact": "alice@test.com",
        "linkedin_url": "https://linkedin.com/in/alice",
        "location": "Paris",
        "years_of_experience": 8,
        "daily_rate": 600,
        "annual_salary": None,
        "availability_status": None,
        "availability_date": None,
        "work_mode": None,
        "location_preference": None,
        "mission_duration": None,
        "contract_type": None,
        "preferred_domains": None,
    }
    profile = MagicMock()
    for k, v in {**defaults, **kwargs}.items():
        setattr(profile, k, v)
    return profile


def _mock_exp(**kwargs: object) -> MagicMock:
    defaults: dict[str, object] = {
        "client_name": "TechCorp",
        "role": "Backend Developer",
        "start_date": date(2022, 1, 1),
        "end_date": None,
        "is_current": True,
        "description": "Developed REST APIs",
        "context": "Greenfield project",
        "achievements_summary": "Reduced latency by 30%",
        "skill_usages": [],
        "achievements": [],
    }
    exp = MagicMock()
    for k, v in {**defaults, **kwargs}.items():
        setattr(exp, k, v)
    return exp


def _mock_skill_ref(name: str, kind: SkillKind) -> MagicMock:
    ref = MagicMock()
    ref.name = name
    ref.kind = kind
    return ref


def _mock_skill(name: str, kind: SkillKind, level: str = "", featured: bool = False) -> MagicMock:
    skill = MagicMock()
    skill.skill_ref = _mock_skill_ref(name, kind)
    skill.self_assessed_level = level
    skill.featured = featured
    return skill


def _mock_usage(name: str, kind: SkillKind) -> MagicMock:
    usage = MagicMock()
    usage.skill_ref = _mock_skill_ref(name, kind)
    usage.intensity = UsageIntensity.primary
    return usage


def _mock_education(**kwargs: object) -> MagicMock:
    defaults: dict[str, object] = {
        "school": "INSA Lyon",
        "degree": "Diplome d'ingenieur",
        "field_of_study": "Informatique",
        "start_date": date(2010, 9, 1),
        "end_date": date(2015, 6, 1),
        "description": None,
    }
    edu = MagicMock()
    for k, v in {**defaults, **kwargs}.items():
        setattr(edu, k, v)
    return edu


def _mock_certification(**kwargs: object) -> MagicMock:
    defaults: dict[str, object] = {
        "name": "AWS Solutions Architect",
        "issuer": "Amazon Web Services",
        "issue_date": date(2022, 5, 1),
        "expiry_date": None,
        "credential_url": None,
    }
    cert = MagicMock()
    for k, v in {**defaults, **kwargs}.items():
        setattr(cert, k, v)
    return cert


def _mock_language(**kwargs: object) -> MagicMock:
    defaults: dict[str, object] = {"name": "Anglais", "level": LanguageLevel.C1}
    language = MagicMock()
    for k, v in {**defaults, **kwargs}.items():
        setattr(language, k, v)
    return language


# ---------------------------------------------------------------------------
# Simple field replacement — templates now use standard field names directly
# ---------------------------------------------------------------------------


def test_simple_placeholder_replaced() -> None:
    path = _make_docx_path(["Nom: {{last_name}}", "Prénom: {{first_name}}"])
    result = generate_document(path, _mock_profile(), [], [])
    doc = Document(io.BytesIO(result))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "Martin" in texts
    assert "Alice" in texts
    assert "{{last_name}}" not in texts


def test_unknown_field_renders_as_empty() -> None:
    """Undefined Jinja2 variables render as empty string."""
    path = _make_docx_path(["Data: {{ghost_field}}"])
    result = generate_document(path, _mock_profile(), [], [])
    doc = Document(io.BytesIO(result))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "{{ghost_field}}" not in texts


# ---------------------------------------------------------------------------
# Experience block — paragraph syntax
# ---------------------------------------------------------------------------


def test_experience_block_repeated_per_item() -> None:
    path = _make_docx_path(
        [
            "{%p for exp in experiences %}",
            "{{exp.client_name}} — {{exp.role}}",
            "{%p endfor %}",
        ]
    )
    exp1 = _mock_exp(client_name="Alpha", role="Dev")
    exp2 = _mock_exp(client_name="Beta", role="Lead")
    result = generate_document(path, _mock_profile(), [exp1, exp2], [])
    doc = Document(io.BytesIO(result))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "Alpha" in texts and "Beta" in texts
    assert "{%p" not in texts and "{{" not in texts


def test_no_experiences_removes_block_markers() -> None:
    path = _make_docx_path(
        [
            "Header",
            "{%p for exp in experiences %}",
            "{{exp.client_name}}",
            "{%p endfor %}",
            "Footer",
        ]
    )
    result = generate_document(path, _mock_profile(), [], [])
    doc = Document(io.BytesIO(result))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "Header" in texts and "Footer" in texts
    assert "{%p" not in texts


def test_experience_current_end_date_shows_present() -> None:
    path = _make_docx_path(
        [
            "{%p for exp in experiences %}",
            "{{exp.start_date}} - {{exp.end_date}}",
            "{%p endfor %}",
        ]
    )
    exp = _mock_exp(start_date=date(2022, 6, 1), end_date=None, is_current=True)
    result = generate_document(path, _mock_profile(), [exp], [])
    doc = Document(io.BytesIO(result))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "06/2022" in texts and "présent" in texts


def test_date_formatted_mm_yyyy() -> None:
    path = _make_docx_path(
        [
            "{%p for exp in experiences %}",
            "{{exp.start_date}} to {{exp.end_date}}",
            "{%p endfor %}",
        ]
    )
    exp = _mock_exp(start_date=date(2021, 3, 15), end_date=date(2023, 11, 1), is_current=False)
    result = generate_document(path, _mock_profile(), [exp], [])
    doc = Document(io.BytesIO(result))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "03/2021" in texts and "11/2023" in texts


def test_achievements_summary_rendered() -> None:
    path = _make_docx_path(
        [
            "{%p for exp in experiences %}",
            "Summary: {{exp.achievements_summary}}",
            "{%p endfor %}",
        ]
    )
    exp = _mock_exp(achievements_summary="Reduced latency by 40%")
    result = generate_document(path, _mock_profile(), [exp], [])
    doc = Document(io.BytesIO(result))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "Reduced latency by 40%" in texts


def test_generate_replaces_annual_salary_placeholder() -> None:
    import tempfile

    doc = Document()
    doc.add_paragraph("Salaire annuel souhaité : {{annual_salary}} €")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        template_path = tmp.name

    profile = _mock_profile(annual_salary=55000)
    result = generate_document(template_path, profile, [], [])
    out_doc = Document(io.BytesIO(result))
    text = "\n".join(p.text for p in out_doc.paragraphs)
    assert "55000" in text
    assert "{{annual_salary}}" not in text


def test_complete_dossier_context_uses_per_experience_skills_and_extra_sections() -> None:
    path = _make_docx_path(
        [
            "Disponibilite: {{availability_label}}",
            "{%p for exp in experiences %}",
            "{{exp.client_name}} stack {{exp.skills_tool|map(attribute='name')|join(', ')}}",
            "{%p endfor %}",
            "{%p for edu in educations %}",
            "Formation {{edu.degree}} {{edu.school}}",
            "{%p endfor %}",
            "{%p for cert in certifications %}",
            "Certification {{cert.name}} {{cert.issuer}}",
            "{%p endfor %}",
            "{%p for lang in languages %}",
            "Langue {{lang.name}} {{lang.level_label}}",
            "{%p endfor %}",
        ]
    )
    profile = _mock_profile(availability_status="available_now")
    exp = _mock_exp(client_name="Alpha", skill_usages=[_mock_usage("Docker", SkillKind.tool)])
    global_skills = [
        _mock_skill("PostgreSQL", SkillKind.tool, "avance", True),
        _mock_skill("Python", SkillKind.technical, "avance", True),
    ]

    result = generate_document(
        path,
        profile,
        [exp],
        global_skills,
        [_mock_education()],
        [_mock_certification()],
        [_mock_language()],
    )

    doc = Document(io.BytesIO(result))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Disponible immédiatement" in text
    assert "Alpha stack Docker" in text
    assert "Alpha stack PostgreSQL" not in text
    assert "Formation Diplome d'ingenieur INSA Lyon" in text
    assert "Certification AWS Solutions Architect Amazon Web Services" in text
    assert "Langue Anglais Avancé" in text


def test_build_context_hides_finances_and_contact_when_scoped_out() -> None:
    from services.documents.docx_engine import build_context

    ctx = build_context(
        profile=_mock_profile(daily_rate=600, annual_salary=55000),
        experiences=[],
        skills=[],
        share_finances=False,
        share_contact=False,
    )
    assert ctx["phone"] == ""
    assert ctx["email_contact"] == ""
    assert ctx["linkedin_url"] == ""
    assert ctx["daily_rate"] == ""
    assert ctx["annual_salary"] == ""


def test_build_context_exposes_known_top_level_keys() -> None:
    from services.documents.docx_engine import build_context

    context = build_context(profile=_mock_profile(), experiences=[], skills=[])
    for key in (
        "first_name",
        "availability_label",
        "work_mode_label",
        "contract_type_label",
        "mission_duration_label",
        "experiences",
        "skills",
        "educations",
        "certifications",
        "languages",
        "featured_achievements",
        "skill_groups",
    ):
        assert key in context
