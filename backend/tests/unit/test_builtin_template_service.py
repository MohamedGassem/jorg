import io

from docx import Document

import services.documents.builtin_template_service as builtin_template_service


def test_builtin_templates_render_mock_previews_without_unresolved_tags() -> None:
    templates = builtin_template_service.list_builtin_templates()

    assert {template.key for template in templates} == {
        "compact_esn",
        "dossier_technique",
        "profil_premium",
    }

    for template in templates:
        rendered = builtin_template_service.render_mock_preview(template)
        doc = Document(io.BytesIO(rendered))
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        text = "\n".join(p.text for p in paragraphs)

        assert "{{" not in text
        assert "{%" not in text
        assert "joris" in text.lower()

        # Formation, certifications and languages render as three separate sections,
        # each heading carrying the accent square "■" of the design system
        for heading in ("formation", "certifications", "langues"):
            assert any(
                p.text.strip().startswith("■") and p.text.strip().lower().endswith(heading)
                for p in paragraphs
            ), f"{template.key}: missing '■ {heading}' section heading"
        assert "formation, certifications" not in text.lower()

        # The commercial band is a table holding availability, daily rate
        # and experience together
        band = next(
            (table for table in doc.tables if "Disponibilité" in table.rows[0].cells[0].text),
            None,
        )
        assert band is not None, f"{template.key}: missing commercial band table"
        band_text = " | ".join(cell.text for row in band.rows for cell in row.cells)
        assert "850 €" in band_text
        assert "9 ans" in band_text
