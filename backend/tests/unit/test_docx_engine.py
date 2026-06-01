# backend/tests/unit/test_docx_engine.py
"""Unit tests for the docxtpl-based generation engine."""

import io
import pathlib
import time
import warnings
from datetime import date
from unittest.mock import MagicMock

import pytest
from docx import Document

from services.docx_engine import (
    exp_flat,
    fmt_date,
    generate_document,
    profile_flat,
    skill_flat,
)


def _mock_profile(**kwargs: object) -> MagicMock:
    profile = MagicMock()
    profile.first_name = kwargs.get("first_name", "Alice")
    profile.last_name = kwargs.get("last_name", "Martin")
    profile.title = kwargs.get("title", "Dev")
    profile.summary = kwargs.get("summary", "")
    profile.phone = kwargs.get("phone", "")
    profile.email_contact = kwargs.get("email_contact", "")
    profile.linkedin_url = kwargs.get("linkedin_url", "")
    profile.location = kwargs.get("location", "")
    profile.years_of_experience = kwargs.get("years_of_experience")
    profile.daily_rate = kwargs.get("daily_rate")
    profile.annual_salary = kwargs.get("annual_salary")
    profile.availability_status = kwargs.get("availability_status")
    profile.work_mode = kwargs.get("work_mode")
    profile.location_preference = kwargs.get("location_preference")
    profile.mission_duration = kwargs.get("mission_duration")
    profile.contract_type = kwargs.get("contract_type")
    profile.preferred_domains = kwargs.get("preferred_domains")
    return profile


def _mock_experience(**kwargs: object) -> MagicMock:
    exp = MagicMock()
    exp.client_name = kwargs.get("client_name", "Acme")
    exp.role = kwargs.get("role", "Engineer")
    exp.start_date = kwargs.get("start_date", date(2022, 1, 1))
    exp.end_date = kwargs.get("end_date", date(2023, 6, 1))
    exp.is_current = kwargs.get("is_current", False)
    exp.description = kwargs.get("description", "desc")
    exp.context = kwargs.get("context", "ctx")
    exp.achievements_summary = kwargs.get("achievements_summary", "ach")
    # technologies removed
    return exp


def _mock_skill(**kwargs: object) -> MagicMock:
    sk = MagicMock()
    skill_ref = MagicMock()
    skill_ref.name = kwargs.get("name", "Python")
    skill_ref.kind = kwargs.get("kind", _FakeEnum("technical"))
    sk.skill_ref = skill_ref
    sk.self_assessed_level = kwargs.get("self_assessed_level", "Expert")
    sk.featured = kwargs.get("featured", False)
    return sk


class _FakeEnum:
    def __init__(self, value: str) -> None:
        self.value = value


# ---------------------------------------------------------------------------
# Pure helper functions
# ---------------------------------------------------------------------------


class TestFmtDate:
    def test_returns_empty_string_for_none(self):
        assert fmt_date(None) == ""

    def test_formats_date_as_mm_yyyy(self):
        assert fmt_date(date(2023, 6, 15)) == "06/2023"


class TestProfileFlat:
    def test_returns_first_name(self):
        p = _mock_profile(first_name="Bob")
        assert profile_flat(p)["first_name"] == "Bob"

    def test_returns_empty_string_for_none_fields(self):
        p = _mock_profile(phone=None)
        assert profile_flat(p)["phone"] == ""

    def test_contains_all_expected_keys(self):
        p = _mock_profile()
        expected = {
            "first_name",
            "last_name",
            "title",
            "summary",
            "phone",
            "email_contact",
            "linkedin_url",
            "location",
            "years_of_experience",
            "daily_rate",
            "annual_salary",
            "availability_status",
            "work_mode",
            "location_preference",
            "mission_duration",
            "contract_type",
            "preferred_domains",
        }
        assert set(profile_flat(p).keys()) == expected

    def test_includes_contract_type(self):
        p = _mock_profile(contract_type=_FakeEnum("freelance"))
        assert profile_flat(p)["contract_type"] == "freelance"

    def test_includes_preferred_domains(self):
        p = _mock_profile(preferred_domains=["finance", "tech"])
        assert profile_flat(p)["preferred_domains"] == "finance, tech"

    def test_preferred_domains_none_returns_empty(self):
        p = _mock_profile(preferred_domains=None)
        assert profile_flat(p)["preferred_domains"] == ""

    def test_includes_annual_salary(self):
        p = _mock_profile(annual_salary=60000)
        assert profile_flat(p)["annual_salary"] == "60000"


class TestExpFlat:
    """exp_flat returns clean keys (no 'experience.' prefix) for use as Jinja2 context."""

    def test_formats_end_date(self):
        exp = _mock_experience(end_date=date(2023, 6, 1), is_current=False)
        assert exp_flat(exp)["end_date"] == "06/2023"

    def test_current_experience_shows_present(self):
        exp = _mock_experience(is_current=True)
        assert exp_flat(exp)["end_date"] == "présent"

    def test_achievements_summary_key_present(self):
        exp = _mock_experience(achievements_summary="Shipped feature X")
        assert exp_flat(exp)["achievements_summary"] == "Shipped feature X"

    def test_achievements_backward_compat_alias(self):
        exp = _mock_experience(achievements_summary="Shipped feature X")
        assert exp_flat(exp)["achievements"] == "Shipped feature X"

    def test_keys_have_no_experience_prefix(self):
        exp = _mock_experience()
        keys = set(exp_flat(exp).keys())
        assert "client_name" in keys
        assert "role" in keys
        assert not any(k.startswith("experience.") for k in keys)

    def test_all_expected_keys_present(self):
        exp = _mock_experience()
        expected = {
            "client_name",
            "role",
            "start_date",
            "end_date",
            "description",
            "context",
            "achievements_summary",
            "achievements",
            "technologies",  # backward-compat alias (always empty)
        }
        assert set(exp_flat(exp).keys()) == expected


class TestSkillFlat:
    """skill_flat returns clean keys (no 'skill.' prefix) for use as Jinja2 context."""

    def test_returns_name(self):
        sk = _mock_skill(name="Python")
        assert skill_flat(sk)["name"] == "Python"

    def test_keys_have_no_skill_prefix(self):
        sk = _mock_skill()
        keys = set(skill_flat(sk).keys())
        assert "name" in keys
        assert not any(k.startswith("skill.") for k in keys)

    def test_all_expected_keys_present(self):
        sk = _mock_skill()
        expected = {
            "name",
            "kind",
            "level",
            "self_assessed_level",
            "featured",
            "category",  # backward-compat alias for kind
            "level_rating",  # backward-compat alias (always empty)
        }
        assert set(skill_flat(sk).keys()) == expected

    def test_kind_extracts_value(self):
        sk = _mock_skill(kind=_FakeEnum("technical"))
        assert skill_flat(sk)["kind"] == "technical"

    def test_level_backward_compat_alias(self):
        sk = _mock_skill(self_assessed_level="Expert")
        assert skill_flat(sk)["level"] == "Expert"

    def test_featured_true_returns_true_string(self):
        sk = _mock_skill(featured=True)
        assert skill_flat(sk)["featured"] == "true"


# ---------------------------------------------------------------------------
# generate_document — simple field replacement
# ---------------------------------------------------------------------------


class TestSimpleFieldReplacement:
    def test_returns_bytes(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{{first_name}}")
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(), [], [], {})
        assert isinstance(result, bytes) and len(result) > 0

    def test_replaces_field_in_paragraph(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{{first_name}} {{last_name}}")
        doc.save(str(tmpl))

        result = generate_document(
            str(tmpl), _mock_profile(first_name="Alice", last_name="Martin"), [], [], {}
        )
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs)
        assert "Alice" in text and "Martin" in text

    def test_replaces_field_in_table_cell(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].paragraphs[0].text = "{{first_name}}"
        table.rows[0].cells[1].paragraphs[0].text = "{{last_name}}"
        doc.save(str(tmpl))

        result = generate_document(
            str(tmpl), _mock_profile(first_name="Alice", last_name="Martin"), [], [], {}
        )
        out = Document(io.BytesIO(result))
        cell_texts = [cell.text for row in out.tables[0].rows for cell in row.cells]
        assert "Alice" in cell_texts
        assert "Martin" in cell_texts

    def test_replaces_field_in_header(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("body")
        doc.sections[0].header.paragraphs[0].text = "{{title}}"
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(title="Architecte"), [], [], {})
        out = Document(io.BytesIO(result))
        header_text = out.sections[0].header.paragraphs[0].text
        assert "Architecte" in header_text

    def test_none_value_renders_as_empty(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("Tél: {{phone}}")
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(phone=None), [], [], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs)
        assert "{{phone}}" not in text
        assert "Tél:" in text

    def test_unknown_variable_renders_as_empty(self, tmp_path):
        """Jinja2 undefined variables render as empty string (undefined mode)."""
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{{nonexistent_field}}")
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(), [], [], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs)
        assert "{{nonexistent_field}}" not in text

    def test_replaces_annual_salary(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("Salaire: {{annual_salary}} €")
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(annual_salary=55000), [], [], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs)
        assert "55000" in text


# ---------------------------------------------------------------------------
# generate_document — fragmented runs
# ---------------------------------------------------------------------------


class TestFragmentedRuns:
    def test_placeholder_split_across_runs_is_replaced(self, tmp_path):
        """Placeholder split over multiple Word runs must still be replaced.

        This simulates what Word does when applying formatting in the middle of
        typing a placeholder, or when pasting text with mixed formatting.
        """
        tmpl = tmp_path / "t.docx"
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{")
        para.add_run("first_name")
        para.add_run("}}")
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(first_name="Alice"), [], [], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs)
        assert "Alice" in text
        assert "{{" not in text

    def test_placeholder_split_in_three_ways_is_replaced(self, tmp_path):
        """Worst-case fragmentation: each character group is a separate run."""
        tmpl = tmp_path / "t.docx"
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{")
        para.add_run("{last")
        para.add_run("_name}}")
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(last_name="Martin"), [], [], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs)
        assert "Martin" in text


# ---------------------------------------------------------------------------
# generate_document — experience block (paragraphs)
# ---------------------------------------------------------------------------


class TestExperienceBlockParagraphs:
    def test_block_clones_per_experience(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{%p for exp in experiences %}")
        doc.add_paragraph("{{exp.role}} — {{exp.client_name}}")
        doc.add_paragraph("{%p endfor %}")
        doc.save(str(tmpl))

        exp1 = _mock_experience(role="Engineer", client_name="Alpha")
        exp2 = _mock_experience(role="Architect", client_name="Beta")
        result = generate_document(str(tmpl), _mock_profile(), [exp1, exp2], [], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs if p.text)
        assert "Engineer" in text and "Alpha" in text
        assert "Architect" in text and "Beta" in text

    def test_empty_list_removes_block_markers(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("Header")
        doc.add_paragraph("{%p for exp in experiences %}")
        doc.add_paragraph("{{exp.role}}")
        doc.add_paragraph("{%p endfor %}")
        doc.add_paragraph("Footer")
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(), [], [], {})
        texts = [p.text for p in Document(io.BytesIO(result)).paragraphs if p.text]
        assert "Header" in texts
        assert "Footer" in texts
        assert not any("{%" in t or "{{" in t for t in texts)

    def test_current_experience_shows_present(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{%p for exp in experiences %}")
        doc.add_paragraph("{{exp.start_date}} — {{exp.end_date}}")
        doc.add_paragraph("{%p endfor %}")
        doc.save(str(tmpl))

        exp = _mock_experience(start_date=date(2022, 6, 1), end_date=None, is_current=True)
        result = generate_document(str(tmpl), _mock_profile(), [exp], [], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs)
        assert "06/2022" in text and "présent" in text

    def test_achievements_summary_rendered(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{%p for exp in experiences %}")
        doc.add_paragraph("Summary: {{exp.achievements_summary}}")
        doc.add_paragraph("{%p endfor %}")
        doc.save(str(tmpl))

        exp = _mock_experience(achievements_summary="Reduced latency by 40%")
        result = generate_document(str(tmpl), _mock_profile(), [exp], [], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs)
        assert "Reduced latency by 40%" in text


# ---------------------------------------------------------------------------
# generate_document — experience block (table rows)
# ---------------------------------------------------------------------------


class TestExperienceBlockTableRows:
    def _make_row_block_template(self, tmp_path: "pathlib.Path", n_cols: int = 2) -> str:
        """Template: table with marker row / content row / endmarker row."""
        tmpl = tmp_path / "t.docx"
        doc = Document()
        table = doc.add_table(rows=3, cols=n_cols)
        table.rows[0].cells[0].paragraphs[0].text = "{%tr for exp in experiences %}"
        table.rows[1].cells[0].paragraphs[0].text = "{{exp.client_name}}"
        table.rows[1].cells[1].paragraphs[0].text = "{{exp.role}}"
        table.rows[2].cells[0].paragraphs[0].text = "{%tr endfor %}"
        doc.save(str(tmpl))
        return str(tmpl)

    def test_one_row_per_experience(self, tmp_path):
        tmpl = self._make_row_block_template(tmp_path)
        exp1 = _mock_experience(client_name="Alpha", role="Dev")
        exp2 = _mock_experience(client_name="Beta", role="Lead")

        result = generate_document(tmpl, _mock_profile(), [exp1, exp2], [], {})
        out = Document(io.BytesIO(result))
        assert len(out.tables) == 1
        assert len(out.tables[0].rows) == 2

    def test_row_content_matches_experiences(self, tmp_path):
        tmpl = self._make_row_block_template(tmp_path)
        exp1 = _mock_experience(client_name="Alpha", role="Dev")
        exp2 = _mock_experience(client_name="Beta", role="Lead")

        result = generate_document(tmpl, _mock_profile(), [exp1, exp2], [], {})
        out = Document(io.BytesIO(result))
        all_text = " ".join(cell.text for row in out.tables[0].rows for cell in row.cells)
        assert "Alpha" in all_text and "Dev" in all_text
        assert "Beta" in all_text and "Lead" in all_text

    def test_empty_list_produces_zero_rows(self, tmp_path):
        tmpl = self._make_row_block_template(tmp_path)
        result = generate_document(tmpl, _mock_profile(), [], [], {})
        out = Document(io.BytesIO(result))
        assert len(out.tables[0].rows) == 0

    def test_no_marker_text_in_output(self, tmp_path):
        tmpl = self._make_row_block_template(tmp_path)
        exp = _mock_experience(client_name="Alpha", role="Dev")

        result = generate_document(tmpl, _mock_profile(), [exp], [], {})
        out = Document(io.BytesIO(result))
        all_text = " ".join(cell.text for row in out.tables[0].rows for cell in row.cells)
        assert "{%tr" not in all_text
        assert "{{" not in all_text


# ---------------------------------------------------------------------------
# generate_document — skills block (table rows)
# ---------------------------------------------------------------------------


class TestSkillsBlockTableRows:
    def test_one_row_per_skill(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].paragraphs[0].text = "{%tr for sk in skills %}"
        table.rows[1].cells[0].paragraphs[0].text = "{{sk.name}}"
        table.rows[1].cells[1].paragraphs[0].text = "{{sk.kind}}"
        table.rows[2].cells[0].paragraphs[0].text = "{%tr endfor %}"
        doc.save(str(tmpl))

        sk1 = _mock_skill(name="Python", kind=_FakeEnum("technical"))
        sk2 = _mock_skill(name="Django", kind=_FakeEnum("technical"))
        result = generate_document(str(tmpl), _mock_profile(), [], [sk1, sk2], {})
        out = Document(io.BytesIO(result))
        assert len(out.tables[0].rows) == 2
        all_text = " ".join(cell.text for row in out.tables[0].rows for cell in row.cells)
        assert "Python" in all_text and "Django" in all_text

    def test_skills_block_in_paragraphs(self, tmp_path):
        """Non-regression: skills block in paragraphs still works."""
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{%p for sk in skills %}")
        doc.add_paragraph("{{sk.name}} ({{sk.kind}})")
        doc.add_paragraph("{%p endfor %}")
        doc.save(str(tmpl))

        sk1 = _mock_skill(name="Python", kind=_FakeEnum("technical"))
        sk2 = _mock_skill(name="Django", kind=_FakeEnum("technical"))
        result = generate_document(str(tmpl), _mock_profile(), [], [sk1, sk2], {})
        text = " ".join(p.text for p in Document(io.BytesIO(result)).paragraphs if p.text)
        assert "Python" in text and "Django" in text
        assert "{%p" not in text

    def test_empty_skills_produces_zero_rows(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        table = doc.add_table(rows=3, cols=1)
        table.rows[0].cells[0].paragraphs[0].text = "{%tr for sk in skills %}"
        table.rows[1].cells[0].paragraphs[0].text = "{{sk.name}}"
        table.rows[2].cells[0].paragraphs[0].text = "{%tr endfor %}"
        doc.save(str(tmpl))

        result = generate_document(str(tmpl), _mock_profile(), [], [], {})
        out = Document(io.BytesIO(result))
        assert len(out.tables[0].rows) == 0


# ---------------------------------------------------------------------------
# generate_document — performance
# ---------------------------------------------------------------------------


class TestPerformance:
    def test_100_experiences_completes_within_2_seconds(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{%p for exp in experiences %}")
        doc.add_paragraph(
            "{{exp.role}} - {{exp.client_name}} ({{exp.start_date}}-{{exp.end_date}})"
        )
        doc.add_paragraph("{%p endfor %}")
        doc.save(str(tmpl))

        exps = [_mock_experience(role=f"Role {i}", client_name=f"Client {i}") for i in range(100)]
        start = time.monotonic()
        result = generate_document(str(tmpl), _mock_profile(), exps, [], {})
        elapsed = time.monotonic() - start

        assert isinstance(result, bytes) and len(result) > 0
        assert elapsed < 2.0, f"Generation took {elapsed:.1f}s — too slow"


# ---------------------------------------------------------------------------
# generate_document — error handling
# ---------------------------------------------------------------------------


class TestErrorHandling:
    def test_missing_template_file_raises_value_error(self):
        with pytest.raises(ValueError, match="unreadable"):
            generate_document("/nonexistent/path/to/template.docx", _mock_profile(), [], [], {})

    def test_corrupt_file_raises_value_error(self, tmp_path):
        bad = tmp_path / "bad.docx"
        bad.write_bytes(b"this is not a zip/docx file at all")
        with pytest.raises(ValueError, match="unreadable"):
            generate_document(str(bad), _mock_profile(), [], [], {})

    def test_invalid_jinja2_syntax_raises_value_error(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{%p for exp in experiences %}")
        doc.add_paragraph("{{exp.role}}")
        # intentionally missing endfor — Jinja2 will raise TemplateSyntaxError
        doc.save(str(tmpl))
        with pytest.raises(ValueError, match=r"[Jj]inja2|syntax|template"):
            generate_document(str(tmpl), _mock_profile(), [], [], {})

    def test_nonempty_mappings_emits_deprecation_warning(self, tmp_path):
        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{{first_name}}")
        doc.save(str(tmpl))
        with pytest.warns(DeprecationWarning, match="mappings"):
            generate_document(str(tmpl), _mock_profile(), [], [], {"{{first_name}}": "first_name"})

    def test_empty_mappings_does_not_warn(self, tmp_path):

        tmpl = tmp_path / "t.docx"
        doc = Document()
        doc.add_paragraph("{{first_name}}")
        doc.save(str(tmpl))
        with warnings.catch_warnings():
            warnings.simplefilter("error", DeprecationWarning)
            # should not raise
            generate_document(str(tmpl), _mock_profile(), [], [], {})


# ---------------------------------------------------------------------------
# _group_skills_by_kind
# ---------------------------------------------------------------------------


from enum import StrEnum  # noqa: E402

from models.skill import SkillKind  # noqa: E402
from services.docx_engine import SkillReferenceProtocol, _group_skills_by_kind  # noqa: E402


class _FakeRef:
    def __init__(self, kind: SkillKind) -> None:
        self.kind: StrEnum | None = kind
        self.name: str | None = "FakeSkill"


class _FakeSkill:
    def __init__(self, kind: SkillKind, featured: bool = False) -> None:
        self.skill_ref: SkillReferenceProtocol = _FakeRef(kind)
        self.featured: bool = featured
        self.self_assessed_level: str | None = None


def test_group_skills_by_kind_returns_all_kind_keys() -> None:
    skills = [_FakeSkill(SkillKind.technical), _FakeSkill(SkillKind.tool)]
    result = _group_skills_by_kind(skills)
    assert "skills_technical" in result
    assert "skills_tool" in result
    assert "skills_functional" in result  # vide mais présent
    assert "skills_featured" in result


def test_group_skills_by_kind_filters_correctly() -> None:
    s1 = _FakeSkill(SkillKind.technical)
    s2 = _FakeSkill(SkillKind.tool)
    s3 = _FakeSkill(SkillKind.technical)
    result = _group_skills_by_kind([s1, s2, s3])
    assert len(result["skills_technical"]) == 2
    assert len(result["skills_tool"]) == 1
    assert len(result["skills_methodology"]) == 0


def test_group_skills_featured_first_within_type() -> None:
    plain = _FakeSkill(SkillKind.technical, featured=False)
    starred = _FakeSkill(SkillKind.technical, featured=True)
    result = _group_skills_by_kind([plain, starred])
    tech = result["skills_technical"]
    assert tech[0]["featured"] == "true"
    assert tech[1]["featured"] == "false"


def test_group_skills_featured_cross_type() -> None:
    s1 = _FakeSkill(SkillKind.technical, featured=True)
    s2 = _FakeSkill(SkillKind.tool, featured=False)
    s3 = _FakeSkill(SkillKind.methodology, featured=True)
    result = _group_skills_by_kind([s1, s2, s3])
    assert len(result["skills_featured"]) == 2
