# backend/services/docx_parser.py
"""Extract {{...}} placeholders from a Word .docx file."""

from __future__ import annotations

import re
from typing import Any

from docx import Document  # type: ignore[import-untyped,unused-ignore]

_PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")

# Old Mustache block markers: {{#EXPERIENCES}}, {{/EXPERIENCES}}
_MUSTACHE_BLOCK_RE = re.compile(r"^\{\{[#/]")

# Jinja2 loop variable access: {{exp.*}}, {{sk.*}}, {{edu.*}}, {{cert.*}}, {{lang.*}}
# These are resolved inside {%p for ... %} / {%tr for ... %} blocks and are
# not standalone mappable fields (see docs/template-syntax.md).
_LOOP_VAR_RE = re.compile(r"^\{\{(exp|sk|edu|cert|lang)\.")


def _iter_paragraphs(doc: Any) -> list[str]:
    """Collect all text blocks from paragraphs, table cells, headers, and footers."""
    texts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if hf is not None:
                texts.extend(p.text for p in hf.paragraphs)
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            texts.append(cell.text)
    return texts


def is_block_marker(placeholder: str) -> bool:
    """Return True if the placeholder is control syntax that should not be mapped.

    Excluded patterns:
    - Old Mustache block markers: ``{{#EXPERIENCES}}``, ``{{/EXPERIENCES}}``
    - Jinja2 loop variable access: ``{{exp.role}}``, ``{{sk.name}}``,
      ``{{edu.degree}}``, ``{{cert.name}}``, ``{{lang.name}}``, etc.
      These appear inside ``{%p for ... %}`` blocks and are resolved
      automatically — recruiters do not map them individually.

    Note: Jinja2 block tags (``{%p for ... %}``, ``{%tr endfor %}``) use ``{%``
    rather than ``{{`` and are never matched by the placeholder regex at all.
    """
    return bool(_MUSTACHE_BLOCK_RE.match(placeholder) or _LOOP_VAR_RE.match(placeholder))


def extract_placeholders(file_path: str) -> list[str]:
    """Return deduplicated list of mappable ``{{...}}`` placeholders in the document.

    Preserves first-occurrence order. Excludes:
    - Old Mustache block markers: ``{{#EXPERIENCES}}``, ``{{/EXPERIENCES}}``
    - Jinja2 loop variables: ``{{exp.*}}``, ``{{sk.*}}``, ``{{edu.*}}``,
      ``{{cert.*}}``, ``{{lang.*}}``
    - Jinja2 block tags (``{%p ... %}``, ``{%tr ... %}``): excluded naturally
      because they use ``{%`` rather than ``{{``.

    What remains are the standalone profile-field placeholders such as
    ``{{first_name}}``, ``{{daily_rate}}``, etc., which recruiters can inspect
    to verify template coverage.
    """
    doc = Document(file_path)
    seen: dict[str, None] = {}
    for text in _iter_paragraphs(doc):
        for match in _PLACEHOLDER_RE.finditer(text):
            ph = match.group()
            if is_block_marker(ph):
                continue
            seen.setdefault(ph, None)
    return list(seen.keys())
