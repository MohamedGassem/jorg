# backend/services/documents/templatize_ops.py
"""Templatize operation plan (LLM output schema) and its deterministic applicator."""

from __future__ import annotations

import copy
import io
from dataclasses import dataclass, field
from typing import Annotated, Any, Literal

from docx import Document  # type: ignore[import-untyped,unused-ignore]
from docx.oxml import OxmlElement  # type: ignore[import-untyped,unused-ignore]
from docx.table import _Row  # type: ignore[import-untyped,unused-ignore]
from docx.text.paragraph import Paragraph  # type: ignore[import-untyped,unused-ignore]
from pydantic import BaseModel, Field


class ParagraphTarget(BaseModel):
    kind: Literal["paragraph"]
    paragraph: int


class CellTarget(BaseModel):
    kind: Literal["cell"]
    table: int
    row: int
    cell: int


class ParagraphsRangeTarget(BaseModel):
    kind: Literal["paragraphs"]
    start_paragraph: int
    end_paragraph: int


class TableRowsRangeTarget(BaseModel):
    kind: Literal["table_rows"]
    table: int
    start_row: int
    end_row: int


TextTarget = Annotated[ParagraphTarget | CellTarget, Field(discriminator="kind")]
BlockTarget = Annotated[ParagraphsRangeTarget | TableRowsRangeTarget, Field(discriminator="kind")]


class ReplaceTextOp(BaseModel):
    op: Literal["replace_text"]
    target: TextTarget
    find: str
    placeholder: str


class WrapParagraphsLoopOp(BaseModel):
    op: Literal["wrap_paragraphs_loop"]
    start_paragraph: int
    end_paragraph: int
    loop_var: str
    collection: str


class WrapTableRowsLoopOp(BaseModel):
    op: Literal["wrap_table_rows_loop"]
    table: int
    start_row: int
    end_row: int
    loop_var: str
    collection: str


class DeleteBlockOp(BaseModel):
    op: Literal["delete_block"]
    target: BlockTarget


class FlagResidualOp(BaseModel):
    op: Literal["flag_residual"]
    target: TextTarget
    note: str


Operation = Annotated[
    ReplaceTextOp | WrapParagraphsLoopOp | WrapTableRowsLoopOp | DeleteBlockOp | FlagResidualOp,
    Field(discriminator="op"),
]


class TemplatizePlan(BaseModel):
    operations: list[Operation]


@dataclass
class ApplyResult:
    docx_bytes: bytes
    rejected: list[str] = field(default_factory=list)
    residual_flags: list[str] = field(default_factory=list)


def _replace_in_paragraph(paragraph: Any, find: str, replacement: str) -> bool:
    """Replace the first occurrence of ``find`` in the paragraph, run-aware.

    If ``find`` is inside a single run, only that run changes. If it spans
    several runs, the replacement goes into the first overlapping run and the
    overlapping remainder is cleared, preserving the first run's formatting.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    start = full_text.find(find)
    if start == -1:
        return False
    end = start + len(find)

    pos = 0
    replaced = False
    for run in paragraph.runs:
        run_start, run_end = pos, pos + len(run.text)
        pos = run_end
        if run_end <= start or run_start >= end:
            continue  # run entirely outside the match
        local_start = max(start - run_start, 0)
        local_end = min(end - run_start, len(run.text))
        if not replaced:
            run.text = run.text[:local_start] + replacement + run.text[local_end:]
            replaced = True
        else:
            run.text = run.text[:local_start] + run.text[local_end:]
    return True


def _insert_paragraph_after(paragraph: Any, text: str) -> None:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    Paragraph(new_p, paragraph._parent).add_run(text)


def _insert_paragraph_before(paragraph: Any, text: str) -> None:
    paragraph.insert_paragraph_before(text)


def _insert_loop_row(table: Any, ref_row: Any, tag: str, after: bool) -> None:
    """Insert a dedicated row containing only ``tag`` in its first cell."""
    new_tr = copy.deepcopy(ref_row._tr)
    if after:
        ref_row._tr.addnext(new_tr)
    else:
        ref_row._tr.addprevious(new_tr)
    new_row = _Row(new_tr, table)
    for cell in new_row.cells:
        cell.text = ""
    new_row.cells[0].text = tag


def apply_operations(file_path: str, plan: TemplatizePlan) -> ApplyResult:
    """Apply the plan to the document. Invalid targets are rejected, never raised."""
    doc = Document(file_path)
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)
    # Freeze rows up front, like paragraphs. Operations insert/delete elements;
    # resolving targets against this snapshot keeps every index aligned with the
    # structure the plan was computed against (live re-reads would shift after a
    # wrap_table_rows_loop and silently mis-target later ops on the same table).
    table_rows = [list(table.rows) for table in tables]
    rejected: list[str] = []
    residual_flags: list[str] = []

    def paragraph_at(index: int) -> Any | None:
        return paragraphs[index] if 0 <= index < len(paragraphs) else None

    def row_at(table_index: int, row_index: int) -> tuple[Any, Any] | None:
        if not 0 <= table_index < len(tables):
            return None
        rows = table_rows[table_index]
        if not 0 <= row_index < len(rows):
            return None
        return tables[table_index], rows[row_index]

    for operation in plan.operations:
        label = operation.op
        if isinstance(operation, ReplaceTextOp):
            if operation.target.kind == "paragraph":
                paragraph = paragraph_at(operation.target.paragraph)
                if paragraph is None or not _replace_in_paragraph(
                    paragraph, operation.find, operation.placeholder
                ):
                    rejected.append(f"{label}: cible ou texte introuvable ({operation.find!r})")
            else:
                located = row_at(operation.target.table, operation.target.row)
                if located is None:
                    rejected.append(f"{label}: tableau/ligne introuvable")
                    continue
                _table, row = located
                cells = row.cells
                if not 0 <= operation.target.cell < len(cells):
                    rejected.append(f"{label}: cellule introuvable")
                    continue
                cell = cells[operation.target.cell]
                if not any(
                    _replace_in_paragraph(p, operation.find, operation.placeholder)
                    for p in cell.paragraphs
                ):
                    rejected.append(f"{label}: texte introuvable ({operation.find!r})")
        elif isinstance(operation, WrapParagraphsLoopOp):
            start = paragraph_at(operation.start_paragraph)
            end = paragraph_at(operation.end_paragraph)
            if start is None or end is None or operation.end_paragraph < operation.start_paragraph:
                rejected.append(f"{label}: bornes de paragraphes invalides")
                continue
            _insert_paragraph_before(
                start, f"{{%p for {operation.loop_var} in {operation.collection} %}}"
            )
            _insert_paragraph_after(end, "{%p endfor %}")
        elif isinstance(operation, WrapTableRowsLoopOp):
            start_located = row_at(operation.table, operation.start_row)
            end_located = row_at(operation.table, operation.end_row)
            if (
                start_located is None
                or end_located is None
                or operation.end_row < operation.start_row
            ):
                rejected.append(f"{label}: bornes de lignes invalides")
                continue
            table, start_row = start_located
            _table, end_row = end_located
            _insert_loop_row(
                table,
                start_row,
                f"{{%tr for {operation.loop_var} in {operation.collection} %}}",
                after=False,
            )
            _insert_loop_row(table, end_row, "{%tr endfor %}", after=True)
        elif isinstance(operation, DeleteBlockOp):
            if operation.target.kind == "paragraphs":
                if operation.target.end_paragraph < operation.target.start_paragraph:
                    rejected.append(f"{label}: bornes invalides")
                    continue
                victims = [
                    paragraph_at(i)
                    for i in range(
                        operation.target.start_paragraph, operation.target.end_paragraph + 1
                    )
                ]
                if any(v is None for v in victims):
                    rejected.append(f"{label}: paragraphe hors bornes")
                    continue
                for victim in victims:
                    assert victim is not None
                    victim._p.getparent().remove(victim._p)
            else:
                if operation.target.end_row < operation.target.start_row:
                    rejected.append(f"{label}: bornes invalides")
                    continue
                located_rows = [
                    row_at(operation.target.table, i)
                    for i in range(operation.target.start_row, operation.target.end_row + 1)
                ]
                if any(r is None for r in located_rows):
                    rejected.append(f"{label}: ligne hors bornes")
                    continue
                for located_row in located_rows:
                    assert located_row is not None
                    _table, row = located_row
                    row._tr.getparent().remove(row._tr)
        elif isinstance(operation, FlagResidualOp):
            residual_flags.append(operation.note)

    buf = io.BytesIO()
    doc.save(buf)
    return ApplyResult(docx_bytes=buf.getvalue(), rejected=rejected, residual_flags=residual_flags)
