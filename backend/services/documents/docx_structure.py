# backend/services/documents/docx_structure.py
"""Indexed, text-only representation of a docx body (sent to the LLM)."""

from __future__ import annotations

from typing import Any

from docx import Document  # type: ignore[import-untyped,unused-ignore]


def extract_structure(file_path: str) -> dict[str, Any]:
    """Return {'paragraphs': [{index, text}], 'tables': [{index, rows: [{index, cells}]}]}.

    Indexes match python-docx body collections (doc.paragraphs / doc.tables),
    which is the addressing scheme used by templatize operations.
    """
    doc = Document(file_path)
    paragraphs = [{"index": i, "text": p.text} for i, p in enumerate(doc.paragraphs)]
    tables = [
        {
            "index": t_idx,
            "rows": [
                {"index": r_idx, "cells": [cell.text for cell in row.cells]}
                for r_idx, row in enumerate(table.rows)
            ],
        }
        for t_idx, table in enumerate(doc.tables)
    ]
    return {"paragraphs": paragraphs, "tables": tables}
