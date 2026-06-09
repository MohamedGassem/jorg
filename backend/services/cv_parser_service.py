"""CV extraction pipeline for candidate-reviewed profile proposals.

The CV is treated as an input hint, not as the source of truth. This service
extracts text, builds a structured proposal, matches known skills, and stores a
pending review record. It never mutates the candidate profile itself.
"""

from __future__ import annotations

import hashlib
import io
import json
import re
from dataclasses import dataclass
from typing import ClassVar, Protocol, TypedDict
from uuid import UUID

from pydantic import BaseModel, Field, ValidationError
from sqlalchemy import select
from sqlalchemy.exc import ProgrammingError
from sqlalchemy.ext.asyncio import AsyncSession

from models.candidate_profile import CVExtractionProposal, CVExtractionStatus
from models.skill import SkillKind, SkillReference
from services.language_reference_service import LanguageIndex
from services.text_utils import normalize_text as _normalise

MAX_CV_BYTES = 5 * 1024 * 1024
MIN_USABLE_TEXT_CHARS = 180
_MIN_SKILL_LEN = 3
_MAX_NGRAM_WORDS = 5
_MAX_SKILL_SUGGESTIONS = 60

ExtractionMethod = str


class UnsupportedCVFormatError(Exception):
    """Raised when the uploaded file type is not supported."""


class CVTooLargeError(Exception):
    """Raised when the uploaded file exceeds MAX_CV_BYTES."""


class CVTextExtractionError(Exception):
    """Raised when no readable text could be extracted from the file."""


class CVLLMExtractionError(Exception):
    """Raised when an LLM response cannot be validated as strict JSON."""


class CVPersistenceUnavailableError(Exception):
    """Raised when the proposal store is unavailable or not migrated."""


_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(r"(?:(?:\+|00)\d{1,3}[\s.\-]?)?(?:\(?\d{1,4}\)?[\s.\-]?){2,5}\d{2,4}")
_LINKEDIN_RE = re.compile(
    r"(?:https?://)?(?:[a-z]{2,3}\.)?linkedin\.com/in/[A-Za-z0-9_%\-]+/?",
    re.IGNORECASE,
)
_DATE_RE = re.compile(
    r"\b(?:19|20)\d{2}\b|\b\d{1,2}[/-]\d{4}\b|\b(?:janvier|février|mars|avril|mai|juin|"
    r"juillet|août|septembre|octobre|novembre|décembre|january|february|march|april|may|"
    r"june|july|august|september|october|november|december)\s+(?:19|20)\d{2}\b",
    re.IGNORECASE,
)
_SECTION_KEYWORDS = {
    "experience": (
        "expérience",
        "experiences",
        "experience",
        "work experience",
        "employment",
        "experience professionnelle",
        "experiences professionnelles",
    ),
    "education": (
        "formation",
        "formations",
        "education",
        "éducation",
        "diplômes",
        "diplomes",
    ),
    "skills": ("compétences", "competences", "skills", "technologies", "outils"),
    "languages": ("langues", "langages", "languages"),
    "certifications": ("certifications", "certification"),
    "interests": (
        "centres d'intérêt",
        "centres d\u2019intérêt",
        "loisirs",
        "hobbies",
        "interests",
    ),
}
_SECTION_KEYWORDS_NORMALIZED: dict[str, frozenset[str]] = {
    section: frozenset(_normalise(k) for k in keywords)
    for section, keywords in _SECTION_KEYWORDS.items()
}
_GENERIC_SKILL_PREFIXES = {
    "applications deploiement",
    "applications et deploiement",
    "centres d interet",
    "competences",
    "langages",
    "languages",
    "langues",
    "outils",
    "technologies",
}
_HUMAN_LANGUAGES = {
    "anglais": "Anglais",
    "english": "Anglais",
    "espagnol": "Espagnol",
    "spanish": "Espagnol",
    "francais": "Français",
    "french": "Français",
    "allemand": "Allemand",
    "german": "Allemand",
    "italien": "Italien",
    "italian": "Italien",
    "portugais": "Portugais",
    "portuguese": "Portugais",
}
_PROGRAMMING_LANGUAGE_NAMES = {
    "bash",
    "c",
    "c plus plus",
    "c sharp",
    "c++",
    "c#",
    "go",
    "java",
    "javascript",
    "kotlin",
    "php",
    "python",
    "r",
    "ruby",
    "rust",
    "scala",
    "sql",
    "typescript",
}
_MONTHS = {
    "jan": "01",
    "janv": "01",
    "janvier": "01",
    "january": "01",
    "feb": "02",
    "fev": "02",
    "fevr": "02",
    "fevrier": "02",
    "february": "02",
    "mar": "03",
    "mars": "03",
    "march": "03",
    "apr": "04",
    "avr": "04",
    "avril": "04",
    "april": "04",
    "mai": "05",
    "may": "05",
    "jun": "06",
    "juin": "06",
    "june": "06",
    "jul": "07",
    "juil": "07",
    "juillet": "07",
    "july": "07",
    "aug": "08",
    "aout": "08",
    "august": "08",
    "sep": "09",
    "sept": "09",
    "septembre": "09",
    "september": "09",
    "oct": "10",
    "octobre": "10",
    "october": "10",
    "nov": "11",
    "novembre": "11",
    "november": "11",
    "dec": "12",
    "decembre": "12",
    "december": "12",
}
_DATE_RANGE_RE = re.compile(
    r"(?P<start>(?:[^\W\d_]+\.?\s+)?(?:19|20)\d{2})\s*[-\u2013\u2014]\s*"
    r"(?P<end>actuel|present|présent|(?:[^\W\d_]+\.?\s+)?(?:19|20)\d{2})",
    re.IGNORECASE,
)
_KNOWN_LOCATION_TOKENS: frozenset[str] = frozenset(
    {
        "lyon fr",
        "angers",
        "nancy",
        "le creusot",
        "sarreguemines",
    }
)
_BULLET_CHARS = "•‣▪·◦●○∙*-\u2013\u2014"
# A bullet is a leading glyph (optionally followed by space) or a lone "o"
# sub-bullet that must be followed by whitespace, to avoid matching words.
_BULLET_PREFIX_RE = re.compile(rf"^\s*(?:[{re.escape(_BULLET_CHARS)}]+\s*|o\s+)")


class ExtractedField(BaseModel):
    value: str | None = None
    confidence: float = Field(ge=0, le=1, default=0)
    evidence_text: str | None = None
    source_section: str | None = None
    needs_review: bool = True


class IdentityProposal(BaseModel):
    first_name: ExtractedField = Field(default_factory=ExtractedField)
    last_name: ExtractedField = Field(default_factory=ExtractedField)
    email: ExtractedField = Field(default_factory=ExtractedField)
    phone: ExtractedField = Field(default_factory=ExtractedField)
    linkedin_url: ExtractedField = Field(default_factory=ExtractedField)
    title: ExtractedField = Field(default_factory=ExtractedField)
    location: ExtractedField = Field(default_factory=ExtractedField)


class ExperienceProposal(BaseModel):
    role: ExtractedField = Field(default_factory=ExtractedField)
    client_name: ExtractedField = Field(default_factory=ExtractedField)
    start_date: ExtractedField = Field(default_factory=ExtractedField)
    end_date: ExtractedField = Field(default_factory=ExtractedField)
    # Authoritative "ongoing role" flag derived from the CV ("actuel"/"présent").
    # The front must not re-infer this from an empty end_date, which also matches
    # a finished role whose end date simply failed to parse.
    is_current: bool = False
    description: ExtractedField = Field(default_factory=ExtractedField)
    achievements: list[ExtractedField] = []


class EducationProposal(BaseModel):
    school: ExtractedField = Field(default_factory=ExtractedField)
    degree: ExtractedField = Field(default_factory=ExtractedField)
    field_of_study: ExtractedField = Field(default_factory=ExtractedField)
    start_date: ExtractedField = Field(default_factory=ExtractedField)
    end_date: ExtractedField = Field(default_factory=ExtractedField)


class CertificationProposal(BaseModel):
    name: ExtractedField = Field(default_factory=ExtractedField)
    issuer: ExtractedField = Field(default_factory=ExtractedField)
    issue_date: ExtractedField = Field(default_factory=ExtractedField)
    expiry_date: ExtractedField = Field(default_factory=ExtractedField)


class LanguageProposal(BaseModel):
    name: ExtractedField = Field(default_factory=ExtractedField)
    level: ExtractedField = Field(default_factory=ExtractedField)


class SkillProposal(BaseModel):
    original_label: str
    normalized_label: str | None = None
    match_type: str = "unmatched"
    skill_ref_id: UUID | None = None
    name: str | None = None
    kind: SkillKind | None = None
    confidence: float = Field(ge=0, le=1, default=0.7)
    evidence_text: str | None = None
    source_section: str | None = None
    needs_review: bool = True


class ExtractionMetadata(BaseModel):
    filename: str
    file_hash: str
    extraction_method: str
    quality_score: int
    quality_details: dict[str, int | float | bool]
    parser_warnings: list[str] = []
    status: CVExtractionStatus = CVExtractionStatus.PENDING_REVIEW


class CVStructuredProposal(BaseModel):
    identity: IdentityProposal = Field(default_factory=IdentityProposal)
    experiences: list[ExperienceProposal] = []
    education: list[EducationProposal] = []
    certifications: list[CertificationProposal] = []
    languages: list[LanguageProposal] = []
    skills: list[SkillProposal] = []
    warnings: list[str] = []
    extraction_metadata: ExtractionMetadata | None = None


class QualityScore(BaseModel):
    score: int
    details: dict[str, int | float | bool]
    warnings: list[str] = []


class TextExtractionResult(BaseModel):
    text: str
    method: str
    warnings: list[str] = []
    lines: list[DocumentLine] = []
    quality: QualityScore | None = None


class DocumentLine(BaseModel):
    text: str
    page: int = 1
    x0: float | None = None
    y0: float | None = None
    x1: float | None = None
    y1: float | None = None
    font_size: float | None = None
    is_bold: bool | None = None
    line_index: int


@dataclass(frozen=True)
class SectionBlock:
    name: str
    title: DocumentLine | None
    lines: list[DocumentLine]
    confidence: float


@dataclass(frozen=True)
class DateRange:
    start: str | None
    end: str | None
    is_current: bool
    evidence: str
    confidence: float


class DocumentParser(Protocol):
    method: str

    def supports(self, filename: str) -> bool: ...
    def extract(self, data: bytes) -> str: ...


class FastDocxParser:
    method = "docx_fast"

    def supports(self, filename: str) -> bool:
        return _extension(filename) == "docx"

    def extract(self, data: bytes) -> str:
        return _extract_docx(data)


class FastPdfParser:
    method = "pdf_pymupdf"

    def supports(self, filename: str) -> bool:
        return _extension(filename) == "pdf"

    def extract(self, data: bytes) -> str:
        text = _extract_pdf_layout(data)
        return text if text.strip() else _extract_pdf_legacy(data)


class CVLLMClient(Protocol):
    def extract_profile_json(self, text: str) -> str | None: ...


class NoopCVLLMClient:
    def extract_profile_json(self, text: str) -> str | None:
        return None


@dataclass(frozen=True)
class SkillEntry:
    id: UUID
    name: str
    kind: SkillKind


SkillIndex = dict[str, SkillEntry]


class CVContact(TypedDict):
    email: str | None
    phone: str | None
    linkedin_url: str | None


class CVParseData(CVContact):
    skills: list[SkillEntry]


def validate_cv_file(filename: str, data: bytes) -> None:
    if len(data) > MAX_CV_BYTES:
        raise CVTooLargeError(f"file exceeds {MAX_CV_BYTES} bytes")
    if _extension(filename) not in {"pdf", "docx"}:
        raise UnsupportedCVFormatError("Format non supporté. Utilisez un fichier PDF ou DOCX.")


def extract_text(filename: str, data: bytes) -> str:
    result = extract_text_with_metadata(filename, data)
    return result.text


def extract_text_with_metadata(
    filename: str,
    data: bytes,
    parsers: list[DocumentParser] | None = None,
    fallback_parser: DocumentParser | None = None,
) -> TextExtractionResult:
    validate_cv_file(filename, data)
    warnings: list[str] = []
    selected = next(
        (p for p in (parsers or [FastDocxParser(), FastPdfParser()]) if p.supports(filename)),
        None,
    )
    if selected is None:
        raise UnsupportedCVFormatError("Format non supporté. Utilisez un fichier PDF ou DOCX.")

    try:
        text = selected.extract(data)
        fast_error: Exception | None = None
    except Exception as exc:
        warnings.append("L'extraction rapide a échoué.")
        text = ""
        fast_error = exc

    quality = score_text_quality(text)
    if quality.score < 45 and fallback_parser is not None and fallback_parser.supports(filename):
        try:
            fallback_text = fallback_parser.extract(data)
            fallback_quality = score_text_quality(fallback_text)
            if fallback_quality.score > quality.score:
                return TextExtractionResult(
                    text=clean_extracted_text(fallback_text),
                    method=fallback_parser.method,
                    warnings=[*warnings, *fallback_quality.warnings],
                    quality=fallback_quality,
                )
        except CVTextExtractionError as exc:
            warnings.append(str(exc))

    text = clean_extracted_text(text)
    if not text.strip():
        raise CVTextExtractionError("Aucun texte exploitable n'a pu être extrait du fichier.")
    if fast_error is not None:
        warnings.append(str(fast_error))
    lines = _extract_document_lines(filename, data, text)
    return TextExtractionResult(
        text=text,
        method=selected.method,
        warnings=warnings,
        lines=lines,
        quality=quality,
    )


def _extract_pdf_layout(data: bytes) -> str:
    import fitz  # type: ignore[import-untyped]

    document = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for page in document:
        blocks = page.get_text("blocks", sort=True)
        lines: list[str] = []
        for block in blocks:
            if len(block) < 5:
                continue
            text = str(block[4]).strip()
            if not text:
                continue
            lines.extend(line.strip() for line in text.splitlines() if line.strip())
        pages.append("\n".join(lines))
    return "\n\n".join(page for page in pages if page.strip())


def _extract_document_lines(filename: str, data: bytes, fallback_text: str) -> list[DocumentLine]:
    if _extension(filename) == "pdf":
        try:
            lines = _extract_pdf_document_lines(data)
            if lines:
                return lines
        except Exception:
            pass
    return _text_to_document_lines(fallback_text)


def _extract_pdf_document_lines(data: bytes) -> list[DocumentLine]:
    import fitz

    document = fitz.open(stream=data, filetype="pdf")
    lines: list[DocumentLine] = []
    line_index = 0
    for page_index, page in enumerate(document, start=1):
        payload = page.get_text("dict", sort=True)
        for block in payload.get("blocks", []):
            if block.get("type") != 0:
                continue
            for raw_line in block.get("lines", []):
                spans = raw_line.get("spans", [])
                text = clean_extracted_text(" ".join(span.get("text", "") for span in spans))
                if not text:
                    continue
                bbox = raw_line.get("bbox", [None, None, None, None])
                sizes = [float(span.get("size", 0)) for span in spans if span.get("size")]
                font_names = " ".join(str(span.get("font", "")) for span in spans).lower()
                lines.append(
                    DocumentLine(
                        text=text,
                        page=page_index,
                        x0=bbox[0],
                        y0=bbox[1],
                        x1=bbox[2],
                        y1=bbox[3],
                        font_size=max(sizes) if sizes else None,
                        is_bold="bold" in font_names or "black" in font_names,
                        line_index=line_index,
                    )
                )
                line_index += 1
    return lines


def _text_to_document_lines(text: str) -> list[DocumentLine]:
    return [
        DocumentLine(text=line.strip(), line_index=index)
        for index, line in enumerate(text.splitlines())
        if line.strip()
    ]


def _extract_pdf_legacy(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def score_text_quality(text: str) -> QualityScore:
    punctuation = ".,;:/@+-_()[]'"
    readable = sum(
        1 for c in text if c.isprintable() and (c.isalnum() or c.isspace() or c in punctuation)
    )
    total = max(len(text), 1)
    readable_ratio = readable / total
    line_count = len([line for line in text.splitlines() if line.strip()])
    normalized = _normalise(text)
    section_hits = sum(
        1
        for norm_keywords in _SECTION_KEYWORDS_NORMALIZED.values()
        if any(k in normalized for k in norm_keywords)
    )
    has_email = bool(_EMAIL_RE.search(text))
    has_dates = bool(_DATE_RE.search(text))
    length = len(text.strip())

    score = 0
    score += min(25, int(length / 40))
    score += 15 if has_email else 0
    score += 15 if has_dates else 0
    score += min(25, section_hits * 6)
    score += min(10, line_count)
    score += 10 if readable_ratio >= 0.85 else int(readable_ratio * 10)
    score = min(score, 100)

    warnings: list[str] = []
    if length < MIN_USABLE_TEXT_CHARS:
        warnings.append("Le texte extrait est très court; une relecture attentive est nécessaire.")
    if section_hits < 2:
        warnings.append("Peu de sections de CV ont été reconnues automatiquement.")
    if readable_ratio < 0.75:
        warnings.append("Le texte extrait contient beaucoup de caractères peu lisibles.")

    return QualityScore(
        score=score,
        details={
            "length": length,
            "has_email": has_email,
            "has_dates": has_dates,
            "section_hits": section_hits,
            "readable_ratio": round(readable_ratio, 3),
            "line_count": line_count,
        },
        warnings=warnings,
    )


def clean_extracted_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_contact(text: str) -> CVContact:
    email_match = _EMAIL_RE.search(text)
    linkedin_match = _LINKEDIN_RE.search(text)
    phone: str | None = None
    for candidate in _PHONE_RE.findall(text):
        digits = re.sub(r"\D", "", candidate)
        if 8 <= len(digits) <= 15:
            phone = candidate.strip()
            break

    linkedin_url: str | None = None
    if linkedin_match:
        linkedin_url = linkedin_match.group(0)
        if not linkedin_url.startswith("http"):
            linkedin_url = f"https://{linkedin_url}"

    return CVContact(
        email=email_match.group(0) if email_match else None,
        phone=phone,
        linkedin_url=linkedin_url,
    )


def parse_llm_json_strict(raw_json: str) -> CVStructuredProposal:
    try:
        payload = json.loads(raw_json)
        return CVStructuredProposal.model_validate(payload)
    except (json.JSONDecodeError, ValidationError) as exc:
        raise CVLLMExtractionError(
            "La réponse LLM n'est pas un JSON de proposition valide."
        ) from exc


class SectionDetector:
    aliases: ClassVar[dict[str, frozenset[str]]] = {
        section: frozenset(_normalise(k) for k in keywords)
        for section, keywords in _SECTION_KEYWORDS.items()
    }

    def detect(self, lines: list[DocumentLine]) -> dict[str, SectionBlock]:
        sections: dict[str, SectionBlock] = {
            "identity": SectionBlock("identity", None, [], 0.8),
        }
        current = "identity"
        current_title: DocumentLine | None = None
        scores: dict[str, float] = {"identity": 0.8}

        for line in lines:
            detected, score = self._section_for_line(line)
            if detected is not None:
                current = detected
                current_title = line
                sections.setdefault(current, SectionBlock(current, current_title, [], score))
                scores[current] = max(scores.get(current, 0), score)
                continue
            block = sections.get(current)
            if block is None:
                block = SectionBlock(current, current_title, [], scores.get(current, 0.5))
                sections[current] = block
            block.lines.append(line)

        return {
            name: SectionBlock(name, block.title, block.lines, scores.get(name, block.confidence))
            for name, block in sections.items()
        }

    def _section_for_line(self, line: DocumentLine) -> tuple[str | None, float]:
        normalized = _normalise(line.text)
        if not normalized or len(normalized) > 45:
            return None, 0

        if ":" in line.text:
            prefix, suffix = line.text.split(":", 1)
            prefix_norm = _normalise(prefix)
            if prefix_norm in self.aliases["languages"] and _contains_human_language(suffix):
                return "languages", 0.85
            return None, 0

        for section, aliases in self.aliases.items():
            if normalized in aliases:
                style_bonus = 0.08 if line.is_bold else 0
                size_bonus = 0.05 if (line.font_size or 0) >= 12 else 0
                return section, min(0.98, 0.82 + style_bonus + size_bonus)
        return None, 0


class ExperienceBlockParser:
    def parse(self, lines: list[DocumentLine]) -> list[ExperienceProposal]:
        date_positions: list[tuple[int, DateRange]] = []
        for idx in range(len(lines)):
            date_range = _parse_date_range_at(lines, idx)
            if date_range is not None:
                date_positions.append((idx, date_range))
        experiences: list[ExperienceProposal] = []
        if not date_positions:
            description, achievement_texts = _split_description_and_achievements(
                [line.text for line in lines if len(line.text) > 8]
            )
            if not description and not achievement_texts:
                return []
            return [
                ExperienceProposal(
                    description=_field(description, description, "experience", 0.42, True),
                    achievements=[
                        _field(text, text, "experience", 0.56, True) for text in achievement_texts
                    ],
                )
            ]
        for order, (date_idx, date_range) in enumerate(date_positions):
            prev_date_idx = date_positions[order - 1][0] if order > 0 else -1
            next_date_idx = (
                date_positions[order + 1][0] if order + 1 < len(date_positions) else len(lines)
            )
            before = [
                line
                for line in lines[prev_date_idx + 1 : date_idx]
                if _is_block_header_candidate(line.text)
            ]
            after = lines[date_idx + 1 : next_date_idx]
            if not before and not after:
                continue

            header = before[-1] if before else None
            company_after = next(
                (line for line in after[:3] if _is_block_header_candidate(line.text)),
                None,
            )
            company, role = _parse_experience_header(header.text if header else "")
            description_start = 0
            if company_after is not None and not company:
                company = company_after.text
                description_start = after.index(company_after) + 1

            description, achievement_texts = _split_description_and_achievements(
                [line.text for line in after[description_start:] if _is_description_line(line.text)]
            )
            confidence = 0.78 if (role or company) else 0.55
            role_evidence = header.text if header else None
            company_evidence = company_after.text if company_after is not None else role_evidence
            experiences.append(
                ExperienceProposal(
                    role=_field(role, role_evidence, "experience", confidence, True),
                    client_name=_field(
                        company,
                        company_evidence,
                        "experience",
                        confidence,
                        True,
                    ),
                    start_date=_field(
                        date_range.start,
                        date_range.evidence,
                        "experience",
                        date_range.confidence,
                        True,
                    ),
                    end_date=_field(
                        date_range.end,
                        date_range.evidence,
                        "experience",
                        date_range.confidence,
                        date_range.is_current,
                    ),
                    is_current=date_range.is_current,
                    description=_field(
                        description,
                        description,
                        "experience",
                        0.72 if description else 0,
                        True,
                    ),
                    achievements=[
                        _field(text, text, "experience", 0.68, True) for text in achievement_texts
                    ],
                )
            )
        return experiences


class EducationBlockParser:
    def parse(self, lines: list[DocumentLine]) -> list[EducationProposal]:
        date_positions: list[tuple[int, DateRange]] = []
        for idx in range(len(lines)):
            date_range = _parse_date_range_at(lines, idx)
            if date_range is not None:
                date_positions.append((idx, date_range))
        education: list[EducationProposal] = []
        used: set[int] = set()
        for date_idx, date_range in date_positions:
            previous = [
                line for line in lines[max(0, date_idx - 2) : date_idx] if line.text.strip()
            ]
            following = [line for line in lines[date_idx + 1 : date_idx + 3] if line.text.strip()]
            degree_line = previous[-1] if previous else None
            school_line = (
                following[0] if following else (previous[-2] if len(previous) > 1 else None)
            )
            if school_line is not None and _looks_like_non_education_heading(school_line.text):
                school_line = None
            if degree_line is None and school_line is None:
                continue
            used.add(date_idx)
            parsed_school, degree, field = _parse_education_header(
                degree_line.text if degree_line else ""
            )
            school_value = school_line.text if school_line else parsed_school
            education.append(
                EducationProposal(
                    school=_field(
                        school_value,
                        school_value,
                        "education",
                        0.76,
                        True,
                    ),
                    degree=_field(
                        degree,
                        degree_line.text if degree_line else None,
                        "education",
                        0.76,
                        True,
                    ),
                    field_of_study=_field(
                        field,
                        degree_line.text if degree_line else None,
                        "education",
                        0.7,
                        True,
                    ),
                    start_date=_field(
                        date_range.start,
                        date_range.evidence,
                        "education",
                        0.8,
                        True,
                    ),
                    end_date=_field(date_range.end, date_range.evidence, "education", 0.8, True),
                )
            )

        if education:
            return education

        for idx, line in enumerate(lines[:8]):
            if idx in used or len(line.text) < 5:
                continue
            education.append(
                EducationProposal(school=_field(line.text, line.text, "education", 0.45, True))
            )
        return education[:5]


class LanguageParser:
    def __init__(self, language_index: LanguageIndex | None = None) -> None:
        self._index: LanguageIndex = language_index if language_index else _HUMAN_LANGUAGES

    def parse(self, lines: list[DocumentLine]) -> list[LanguageProposal]:
        languages: dict[str, LanguageProposal] = {}
        for line in lines:
            # Split only on list separators, NOT on " - ": the dash usually joins
            # a language to its level ("Anglais - C1"), which must stay together
            # so _extract_language_level can read the level from the same part.
            for raw_part in re.split(r"[,;]", line.text):
                normalized_part = _normalise(raw_part)
                for normalized_language, display in self._index.items():
                    if re.search(rf"\b{re.escape(normalized_language)}\b", normalized_part):
                        level = _extract_language_level(raw_part)
                        languages[display] = LanguageProposal(
                            name=_field(display, raw_part.strip(), "languages", 0.9, False),
                            level=_field(
                                level,
                                raw_part.strip(),
                                "languages",
                                0.82 if level else 0,
                                level is None,
                            ),
                        )
        return list(languages.values())


class SkillParser:
    def parse(self, lines: list[DocumentLine]) -> list[SkillProposal]:
        labels: list[str] = []
        for line in lines:
            if _contains_human_language(line.text) and _normalise(line.text).startswith("langues"):
                continue
            labels.extend(_skill_labels_from_line(line.text))
        return [
            SkillProposal(
                original_label=label,
                normalized_label=_normalise(label),
                evidence_text=label,
                source_section="skills",
                needs_review=True,
            )
            for label in labels[:80]
            if not _is_human_language_label(label)
        ]


def build_structured_proposal(
    text: str,
    filename: str,
    file_hash: str,
    extraction: TextExtractionResult,
    quality: QualityScore,
    index: SkillIndex,
    llm_client: CVLLMClient | None = None,
    language_index: LanguageIndex | None = None,
) -> CVStructuredProposal:
    warnings = [*extraction.warnings, *quality.warnings]
    proposal: CVStructuredProposal | None = None
    raw_llm = (llm_client or NoopCVLLMClient()).extract_profile_json(text)
    if raw_llm:
        try:
            proposal = parse_llm_json_strict(raw_llm)
        except CVLLMExtractionError:
            warnings.append("Extraction LLM ignorée: JSON invalide.")

    if proposal is None:
        proposal = _deterministic_structured_proposal(text, extraction.lines, language_index)
        warnings.append("Extraction structurée heuristique utilisée; validez chaque champ.")

    proposal.skills = match_structured_skills(text, proposal.skills, index)
    proposal.warnings = [*proposal.warnings, *warnings]
    proposal.extraction_metadata = ExtractionMetadata(
        filename=filename,
        file_hash=file_hash,
        extraction_method=extraction.method,
        quality_score=quality.score,
        quality_details=quality.details,
        parser_warnings=warnings,
    )
    return proposal


def _deterministic_structured_proposal(
    text: str,
    document_lines: list[DocumentLine] | None = None,
    language_index: LanguageIndex | None = None,
) -> CVStructuredProposal:
    contact = extract_contact(text)
    lines = document_lines or _text_to_document_lines(text)
    sections = SectionDetector().detect(lines)
    identity_lines = sections.get("identity", SectionBlock("identity", None, [], 0)).lines
    identity_text = "\n".join(line.text for line in identity_lines)
    identity = IdentityProposal(
        email=_field(contact["email"], contact["email"], "identity", 0.95, False),
        phone=_field(contact["phone"], contact["phone"], "identity", 0.8, True),
        linkedin_url=_field(
            contact["linkedin_url"],
            contact["linkedin_url"],
            "identity",
            0.9,
            False,
        ),
    )
    if identity_lines:
        first_line = identity_lines[0].text
        words = first_line.split()
        if 2 <= len(words) <= 4 and not _EMAIL_RE.search(first_line):
            identity.first_name = _field(words[0], first_line, "identity", 0.7, True)
            identity.last_name = _field(" ".join(words[1:]), first_line, "identity", 0.7, True)
    if len(identity_lines) > 1 and not _EMAIL_RE.search(identity_lines[1].text):
        identity.title = _field(
            identity_lines[1].text,
            identity_lines[1].text,
            "identity",
            0.62,
            True,
        )
    location = _extract_location(identity_text)
    if location:
        identity.location = _field(location, location, "identity", 0.72, True)

    experience_lines = sections.get("experience", SectionBlock("experience", None, [], 0)).lines
    education_lines = sections.get("education", SectionBlock("education", None, [], 0)).lines
    skills_lines = sections.get("skills", SectionBlock("skills", None, [], 0)).lines
    language_lines = [
        *sections.get("languages", SectionBlock("languages", None, [], 0)).lines,
        *identity_lines,
    ]
    certification_lines = sections.get(
        "certifications", SectionBlock("certifications", None, [], 0)
    ).lines

    experiences = ExperienceBlockParser().parse(experience_lines)
    education = EducationBlockParser().parse(education_lines)
    certifications = [
        CertificationProposal(name=_field(line.text, line.text, "certifications", 0.55, True))
        for line in certification_lines[:5]
        if len(line.text) > 5
    ]
    languages = LanguageParser(language_index).parse(language_lines)
    skills = SkillParser().parse(skills_lines)
    return CVStructuredProposal(
        identity=identity,
        experiences=experiences,
        education=education,
        certifications=certifications,
        languages=languages,
        skills=skills,
    )


def _field(
    value: str | None,
    evidence: str | None,
    section: str,
    confidence: float,
    needs_review: bool,
) -> ExtractedField:
    if value is None:
        return ExtractedField()
    return ExtractedField(
        value=value,
        confidence=confidence,
        evidence_text=evidence,
        source_section=section,
        needs_review=needs_review,
    )


def _parse_date_range(text: str) -> DateRange | None:
    match = _DATE_RANGE_RE.search(text)
    if match is None:
        year_match = re.search(
            r"\b((?:19|20)\d{2})\s*[-\u2013\u2014]\s*((?:19|20)\d{2})\b",
            text,
        )
        if not year_match:
            return None
        return DateRange(
            start=year_match.group(1),
            end=year_match.group(2),
            is_current=False,
            evidence=text,
            confidence=0.75,
        )
    end_raw = match.group("end")
    is_current = _normalise(end_raw) in {"actuel", "present"}
    return DateRange(
        start=_normalise_date_value(match.group("start")),
        end=None if is_current else _normalise_date_value(end_raw),
        is_current=is_current,
        evidence=text,
        confidence=0.86 if not is_current else 0.8,
    )


def _parse_date_range_at(lines: list[DocumentLine], index: int) -> DateRange | None:
    direct = _parse_date_range(lines[index].text)
    if direct is not None:
        return direct
    if index + 1 >= len(lines):
        return None
    if not re.search(r"[-\u2013\u2014]", lines[index].text) or not _is_year_only(
        lines[index + 1].text
    ):
        return None
    continuation = f"{lines[index].text} {lines[index + 1].text}"
    return _parse_date_range(continuation)


def _normalise_date_value(value: str) -> str | None:
    cleaned = _normalise(value.replace(".", " "))
    year_match = re.search(r"\b((?:19|20)\d{2})\b", cleaned)
    if not year_match:
        return None
    year = year_match.group(1)
    month = next(
        (number for name, number in _MONTHS.items() if re.search(rf"\b{name}\b", cleaned)),
        None,
    )
    return f"{year}-{month}" if month else year


def _is_block_header_candidate(text: str) -> bool:
    stripped = text.strip()
    if not stripped or _parse_date_range(stripped) is not None or _is_year_only(stripped):
        return False
    if stripped.startswith(("•", "-")):
        return False
    if _EMAIL_RE.search(stripped):
        return False
    if _normalise(stripped) in _KNOWN_LOCATION_TOKENS:
        return False
    return len(stripped) <= 120


def _parse_experience_header(text: str) -> tuple[str | None, str | None]:
    stripped = text.strip()
    if not stripped:
        return None, None
    if "," in stripped:
        company, role = stripped.split(",", 1)
        return company.strip() or None, role.strip() or None
    return None, stripped


def _is_description_line(text: str) -> bool:
    stripped = text.strip()
    if not stripped or _parse_date_range(stripped) is not None or _is_year_only(stripped):
        return False
    if _normalise(stripped) in _KNOWN_LOCATION_TOKENS:
        return False
    return not _is_block_header_candidate(stripped) or stripped.startswith(("•", "-"))


def _is_bullet_line(text: str) -> bool:
    return bool(_BULLET_PREFIX_RE.match(text)) and bool(_strip_bullet(text))


def _strip_bullet(text: str) -> str:
    return _BULLET_PREFIX_RE.sub("", text, count=1).strip()


def _split_description_and_achievements(
    raw_lines: list[str],
) -> tuple[str | None, list[str]]:
    """Split an experience body into a preamble and one achievement per bullet.

    When the block contains typographic bullets, the lines before the first
    bullet become the description and each bulleted line becomes an achievement
    (a non-bulleted line after a bullet is folded into the previous achievement,
    handling wrapped lines). When no bullet glyph survives extraction (e.g. Word
    auto-numbered lists), every line becomes its own achievement and the
    description stays empty, since a preamble cannot be told apart reliably.
    """
    lines = [line for line in raw_lines if line.strip()]
    if not lines:
        return None, []
    bullet_flags = [_is_bullet_line(line) for line in lines]
    if any(bullet_flags):
        first_bullet = bullet_flags.index(True)
        description = " ".join(line.strip() for line in lines[:first_bullet]).strip() or None
        achievements: list[str] = []
        for line, is_bullet in zip(lines[first_bullet:], bullet_flags[first_bullet:], strict=True):
            if is_bullet:
                achievements.append(_strip_bullet(line))
            elif achievements:
                achievements[-1] = f"{achievements[-1]} {line.strip()}".strip()
        return description, [text for text in achievements if text]
    return None, [text for line in lines if (text := _strip_bullet(line))]


def _is_year_only(text: str) -> bool:
    return bool(re.fullmatch(r"(?:19|20)\d{2}", text.strip()))


def _parse_degree_and_field(text: str) -> tuple[str | None, str | None]:
    stripped = text.strip()
    if not stripped or _parse_date_range(stripped):
        return None, None
    if ":" in stripped:
        degree, field = stripped.split(":", 1)
        return degree.strip() or None, field.strip() or None
    return stripped, None


def _parse_education_header(text: str) -> tuple[str | None, str | None, str | None]:
    stripped = text.strip()
    if not stripped:
        return None, None, None
    normalized = _normalise(stripped)
    degree_keywords = ("diplome", "ingenieur", "licence", "bts", "master")
    for keyword in degree_keywords:
        match = re.search(rf"\b{keyword}\b", normalized)
        if match is None:
            continue
        prefix_word_count = len(normalized[: match.start()].split())
        words = stripped.split()
        index = len(" ".join(words[:prefix_word_count]))
        if index > 0:
            index += 1
        if index > 0:
            school = stripped[:index].strip(" ,-")
            degree_text = stripped[index:].strip()
            degree, field = _parse_degree_and_field(degree_text)
            return school or None, degree, field
    degree, field = _parse_degree_and_field(stripped)
    return None, degree, field


def _looks_like_non_education_heading(text: str) -> bool:
    normalized = _normalise(text)
    return normalized in {
        "projets",
        "projects",
        "technologies",
        "competences",
        "skills",
        "centres d interet",
    }


def _contains_human_language(text: str) -> bool:
    normalized = _normalise(text)
    return any(
        re.search(rf"\b{re.escape(language)}\b", normalized) for language in _HUMAN_LANGUAGES
    )


def _is_human_language_label(label: str) -> bool:
    normalized = _normalise(label)
    return normalized in _HUMAN_LANGUAGES and normalized not in _PROGRAMMING_LANGUAGE_NAMES


def _extract_language_level(text: str) -> str | None:
    normalized = _normalise(text)
    if "natif" in normalized or "native" in normalized:
        return "native"
    cefr = re.search(r"\b(A1|A2|B1|B2|C1|C2)\b", text.upper())
    return cefr.group(1) if cefr else None


def _extract_location(text: str) -> str | None:
    for line in text.splitlines():
        if (
            re.search(r"\b(?:Lyon|Paris|Marseille|Toulouse|France)\b", line, re.IGNORECASE)
            and not _EMAIL_RE.search(line)
            and not _PHONE_RE.search(line)
        ):
            return line.strip(" +")
    return None


def _skill_labels_from_line(line: str) -> list[str]:
    candidates: list[str] = []
    if ":" in line:
        prefix, suffix = line.split(":", 1)
        if _normalise(prefix) in _GENERIC_SKILL_PREFIXES:
            candidates.extend(_split_skill_tokens(suffix))
        else:
            candidates.extend(_split_skill_tokens(prefix))
            candidates.extend(_split_skill_tokens(suffix))
    else:
        candidates.extend(_split_skill_tokens(line))

    seen: set[str] = set()
    labels: list[str] = []
    for candidate in candidates:
        label = _clean_skill_label(candidate)
        norm = _normalise(label)
        if not (2 < len(label) <= 80) or norm in seen or norm in _GENERIC_SKILL_PREFIXES:
            continue
        seen.add(norm)
        labels.append(label)
    return labels


def _split_skill_tokens(value: str) -> list[str]:
    parts = re.split(r"[,;•·|/]|\s+(?:&|\+)\s+", value)
    return [part for part in parts if part.strip()]


def _clean_skill_label(value: str) -> str:
    label = value.strip(" -\t()")
    label = re.sub(r"\s+", " ", label)
    return label


def _ngrams(tokens: list[str], max_words: int) -> set[str]:
    grams: set[str] = set()
    n = len(tokens)
    for size in range(1, max_words + 1):
        for i in range(n - size + 1):
            grams.add(" ".join(tokens[i : i + size]))
    return grams


async def build_skill_index(db: AsyncSession) -> SkillIndex:
    result = await db.execute(
        select(SkillReference).where(SkillReference.creator_candidate_id.is_(None))
    )
    index: SkillIndex = {}
    for ref in result.scalars().all():
        entry = SkillEntry(id=ref.id, name=ref.name, kind=ref.kind)
        for phrase in (ref.name, *ref.aliases):
            norm = _normalise(phrase)
            if len(norm) >= _MIN_SKILL_LEN and norm not in index:
                index[norm] = entry
    return index


def match_skills_in_index(text: str, index: SkillIndex) -> list[SkillEntry]:
    tokens = [token for token in _normalise(text).split(" ") if token]
    candidates = _ngrams(tokens, _MAX_NGRAM_WORDS)
    matched: dict[UUID, SkillEntry] = {}
    for gram in candidates:
        hit = index.get(gram)
        if hit is not None:
            matched[hit.id] = hit
    ordered = sorted(matched.values(), key=lambda e: (-len(e.name), e.name.lower()))
    return ordered[:_MAX_SKILL_SUGGESTIONS]


def match_structured_skills(
    text: str,
    proposed_skills: list[SkillProposal],
    index: SkillIndex,
) -> list[SkillProposal]:
    results: dict[UUID | str, SkillProposal] = {}
    for skill in proposed_skills:
        norm = _normalise(skill.original_label)
        hit = index.get(norm)
        if hit is None:
            results[f"unmatched:{norm}:{skill.original_label}"] = skill.model_copy(
                update={"normalized_label": norm, "match_type": "unmatched", "needs_review": True}
            )
            continue
        results[hit.id] = skill.model_copy(
            update={
                "normalized_label": norm,
                "match_type": "explicit",
                "skill_ref_id": hit.id,
                "name": hit.name,
                "kind": hit.kind,
                "confidence": 0.9,
                "needs_review": True,
            }
        )

    for hit in match_skills_in_index(text, index):
        if hit.id not in results:
            results[hit.id] = SkillProposal(
                original_label=hit.name,
                normalized_label=_normalise(hit.name),
                match_type="normalized",
                skill_ref_id=hit.id,
                name=hit.name,
                kind=hit.kind,
                confidence=0.75,
                needs_review=True,
            )
    return list(results.values())[:_MAX_SKILL_SUGGESTIONS]


async def parse_cv(
    filename: str,
    data: bytes,
    index: SkillIndex,
) -> CVParseData:
    text = extract_text(filename, data)
    contact = extract_contact(text)
    skills = match_skills_in_index(text, index)
    return CVParseData(
        email=contact["email"],
        phone=contact["phone"],
        linkedin_url=contact["linkedin_url"],
        skills=skills,
    )


async def parse_and_store_cv_proposal(
    candidate_id: UUID,
    filename: str,
    data: bytes,
    db: AsyncSession,
    index: SkillIndex,
    llm_client: CVLLMClient | None = None,
    fallback_parser: DocumentParser | None = None,
    language_index: LanguageIndex | None = None,
) -> CVExtractionProposal:
    validate_cv_file(filename, data)
    file_hash = hashlib.sha256(data).hexdigest()
    extraction = extract_text_with_metadata(
        filename,
        data,
        fallback_parser=fallback_parser,
    )
    quality = extraction.quality or score_text_quality(extraction.text)
    if quality.score < 20:
        raise CVTextExtractionError("Le texte extrait est trop court ou trop peu lisible.")
    proposal_payload = build_structured_proposal(
        extraction.text,
        filename,
        file_hash,
        extraction,
        quality,
        index,
        llm_client=llm_client,
        language_index=language_index,
    )
    row = CVExtractionProposal(
        candidate_id=candidate_id,
        filename=filename[:255] or "cv",
        file_hash=file_hash,
        raw_text=extraction.text,
        extraction_method=extraction.method,
        quality_score=quality.score,
        quality_details=quality.details,
        proposed_profile=proposal_payload.model_dump(mode="json"),
        warnings=proposal_payload.warnings,
        status=CVExtractionStatus.PENDING_REVIEW,
    )
    db.add(row)
    try:
        await db.commit()
    except ProgrammingError as exc:
        await db.rollback()
        if _is_missing_proposal_table(exc):
            raise CVPersistenceUnavailableError(
                "Le stockage des propositions CV n'est pas prêt. "
                "Appliquez les migrations de base de données."
            ) from exc
        raise
    await db.refresh(row)
    return row


def _extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _is_missing_proposal_table(exc: ProgrammingError) -> bool:
    message = str(exc)
    return "cv_extraction_proposals" in message and (
        "UndefinedTableError" in message or "does not exist" in message
    )
