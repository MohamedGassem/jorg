# backend/services/docx_parser.py
"""Extract {{...}} placeholders from a Word .docx file."""

from __future__ import annotations

import re
from typing import Any

from docx import Document  # type: ignore[import-untyped,unused-ignore]

_PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")
# Block markers start with `#` (opening) or `/` (closing), e.g. {{#EXPERIENCES}}, {{/EXPERIENCES}}
_BLOCK_MARKER_RE = re.compile(r"^\{\{[#/]")


def _iter_paragraphs(doc: Any) -> list[str]:
    """Collect all text blocks from paragraphs and table cells."""
    texts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return texts


def is_block_marker(placeholder: str) -> bool:
    """Return True if the placeholder is a mustache block start/end marker.

    Block markers are control syntax handled by the generator's block expansion
    logic — they must not be presented to recruiters as fields to map.
    """
    return bool(_BLOCK_MARKER_RE.match(placeholder))


def extract_placeholders(file_path: str) -> list[str]:
    """Return deduplicated list of mappable {{...}} placeholders found in the document.

    Preserves first-occurrence order. Excludes block markers such as
    {{#EXPERIENCES}} and {{/EXPERIENCES}} — those are control syntax handled
    by the generator, not data fields.
    """
    doc = Document(file_path)
    seen: dict[str, None] = {}
    for text in _iter_paragraphs(doc):
        for match in _PLACEHOLDER_RE.finditer(text):
            ph = match.group()
            if is_block_marker(ph):
                continue
            seen.setdefault(ph, None)
    return list(seen.keys())
