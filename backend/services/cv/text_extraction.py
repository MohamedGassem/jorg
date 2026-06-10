from __future__ import annotations

import io
import re
from typing import Any, Protocol

from services.cv.constants import MAX_CV_BYTES, _extension
from services.cv.exceptions import CVTextExtractionError, CVTooLargeError, UnsupportedCVFormatError
from services.cv.quality import score_text_quality
from services.cv.schemas import DocumentLine, TextExtractionResult


class DocumentParser(Protocol):
    method: str

    def supports(self, filename: str) -> bool: ...
    def extract(self, data: bytes) -> str: ...


class FastDocxParser:
    method = "docx_fast"

    def supports(self, filename: str) -> bool:
        return _extension(filename) == "docx"

    def extract(self, data: bytes) -> str:
        return _extract_docx(data)


class FastPdfParser:
    method = "pdf_pymupdf"

    def supports(self, filename: str) -> bool:
        return _extension(filename) == "pdf"

    def extract(self, data: bytes) -> str:
        text = _extract_pdf_layout(data)
        return text if text.strip() else _extract_pdf_legacy(data)


def validate_cv_file(filename: str, data: bytes) -> None:
    if len(data) > MAX_CV_BYTES:
        raise CVTooLargeError(f"file exceeds {MAX_CV_BYTES} bytes")
    if _extension(filename) not in {"pdf", "docx"}:
        raise UnsupportedCVFormatError("Format non supporté. Utilisez un fichier PDF ou DOCX.")


def extract_text(filename: str, data: bytes) -> str:
    result = extract_text_with_metadata(filename, data)
    return result.text


def extract_text_with_metadata(
    filename: str,
    data: bytes,
    parsers: list[DocumentParser] | None = None,
    fallback_parser: DocumentParser | None = None,
) -> TextExtractionResult:
    validate_cv_file(filename, data)
    warnings: list[str] = []
    selected = next(
        (p for p in (parsers or [FastDocxParser(), FastPdfParser()]) if p.supports(filename)),
        None,
    )
    if selected is None:
        raise UnsupportedCVFormatError("Format non supporté. Utilisez un fichier PDF ou DOCX.")

    try:
        text = selected.extract(data)
        fast_error: Exception | None = None
    except Exception as exc:
        warnings.append("L'extraction rapide a échoué.")
        text = ""
        fast_error = exc

    quality = score_text_quality(text)
    if quality.score < 45 and fallback_parser is not None and fallback_parser.supports(filename):
        try:
            fallback_text = fallback_parser.extract(data)
            fallback_quality = score_text_quality(fallback_text)
            if fallback_quality.score > quality.score:
                return TextExtractionResult(
                    text=clean_extracted_text(fallback_text),
                    method=fallback_parser.method,
                    warnings=[*warnings, *fallback_quality.warnings],
                    quality=fallback_quality,
                )
        except CVTextExtractionError as exc:
            warnings.append(str(exc))

    text = clean_extracted_text(text)
    if not text.strip():
        raise CVTextExtractionError("Aucun texte exploitable n'a pu être extrait du fichier.")
    if fast_error is not None:
        warnings.append(str(fast_error))
    lines = _extract_document_lines(filename, data, text)
    return TextExtractionResult(
        text=text,
        method=selected.method,
        warnings=warnings,
        lines=lines,
        quality=quality,
    )


def _extract_pdf_layout(data: bytes) -> str:
    lines = _extract_pdf_document_lines(data)
    pages: dict[int, list[str]] = {}
    for line in lines:
        pages.setdefault(line.page or 1, []).append(line.text)
    return "\n\n".join("\n".join(texts) for _, texts in sorted(pages.items()))


def _extract_document_lines(filename: str, data: bytes, fallback_text: str) -> list[DocumentLine]:
    if _extension(filename) == "pdf":
        try:
            lines = _extract_pdf_document_lines(data)
            if lines:
                return lines
        except Exception:
            pass
    return _text_to_document_lines(fallback_text)


# Word gaps in justified PDFs (e.g. LaTeX) can be encoded as glyph-position offsets
# with no space character; measured word gaps sit around 0.13-0.15 x font size while
# intra-word kerning stays below 0.05, so 0.1 separates them with margin on both sides.
_CHAR_GAP_SPACE_RATIO = 0.1


def _span_text_from_chars(span: dict[str, Any]) -> str:
    """Rebuild a rawdict span's text, inserting a space at significant glyph gaps."""
    chars = span.get("chars")
    if not chars:
        return str(span.get("text", ""))
    size = float(span.get("size") or 0)
    threshold = size * _CHAR_GAP_SPACE_RATIO if size else 1.0
    parts: list[str] = []
    prev_x1: float | None = None
    for char in chars:
        text = str(char.get("c", ""))
        bbox = char.get("bbox") or (0.0, 0.0, 0.0, 0.0)
        if (
            prev_x1 is not None
            and text
            and not text.isspace()
            and parts
            and not parts[-1].isspace()
            and bbox[0] - prev_x1 > threshold
        ):
            parts.append(" ")
        parts.append(text)
        prev_x1 = bbox[2]
    return "".join(parts)


def _join_spans(spans: list[dict[str, Any]]) -> str:
    """Join PDF spans into a line string, inserting a space wherever there is a visual gap.

    A plain ``" ".join`` adds spaces between every span pair regardless of whether
    the PDF already has a space character there.  This function checks the x-position
    gap between the right edge of span N and the left edge of span N+1: if the gap
    exceeds 1 pt and neither boundary already carries a whitespace character, a space
    is inserted.  This handles the common case where a PDF encodes adjacent words in
    separate spans without an explicit space character.
    """
    if not spans:
        return ""
    parts: list[str] = [spans[0].get("text", "")]
    for i in range(1, len(spans)):
        prev_bbox = spans[i - 1].get("bbox") or (0, 0, 0, 0)
        curr_bbox = spans[i].get("bbox") or (0, 0, 0, 0)
        gap = curr_bbox[0] - prev_bbox[2]
        curr_text = spans[i].get("text", "")
        needs_space = (
            gap > 1.0
            and bool(parts)
            and not parts[-1].endswith(" ")
            and bool(curr_text)
            and not curr_text[0].isspace()
        )
        if needs_space:
            parts.append(" ")
        parts.append(curr_text)
    return clean_extracted_text("".join(parts))


def _extract_pdf_document_lines(data: bytes) -> list[DocumentLine]:
    import fitz  # type: ignore[import-untyped]

    document = fitz.open(stream=data, filetype="pdf")
    lines: list[DocumentLine] = []
    line_index = 0
    for page_index, page in enumerate(document, start=1):
        payload = page.get_text("rawdict", sort=True)
        for block in payload.get("blocks", []):
            if block.get("type") != 0:
                continue
            for raw_line in block.get("lines", []):
                spans = [
                    {**span, "text": _span_text_from_chars(span)}
                    for span in raw_line.get("spans", [])
                ]
                text = _join_spans(spans)
                if not text:
                    continue
                bbox = raw_line.get("bbox", [None, None, None, None])
                sizes = [float(span.get("size", 0)) for span in spans if span.get("size")]
                font_names = " ".join(str(span.get("font", "")) for span in spans).lower()
                lines.append(
                    DocumentLine(
                        text=text,
                        page=page_index,
                        x0=bbox[0],
                        y0=bbox[1],
                        x1=bbox[2],
                        y1=bbox[3],
                        font_size=max(sizes) if sizes else None,
                        is_bold="bold" in font_names or "black" in font_names,
                        line_index=line_index,
                    )
                )
                line_index += 1
    return _order_lines_in_columns(lines)


def _order_lines_in_columns(lines: list[DocumentLine]) -> list[DocumentLine]:
    """Reorder lines page by page so multi-column layouts read column by column.

    Sidebar layouts (contact column next to the main content) otherwise
    interleave both columns in reading order, which breaks section detection.
    Single-column CVs are untouched because their lines span the page width and
    no vertical gutter exists.
    """
    ordered: list[DocumentLine] = []
    pages: dict[int, list[DocumentLine]] = {}
    for line in lines:
        pages.setdefault(line.page or 1, []).append(line)
    for _, page_lines in sorted(pages.items()):
        ordered.extend(_order_page_columns(page_lines))
    for index, line in enumerate(ordered):
        line.line_index = index
    return ordered


_MIN_GUTTER_WIDTH = 3.0
_MIN_COLUMN_LINES = 5


def _order_page_columns(page_lines: list[DocumentLine]) -> list[DocumentLine]:
    if len(page_lines) < 2 * _MIN_COLUMN_LINES:
        return page_lines
    coords: list[tuple[float, float, float]] = []  # (x0, x1, y0) per line
    for line in page_lines:
        if line.x0 is None or line.x1 is None or line.y0 is None:
            return page_lines
        coords.append((line.x0, line.x1, line.y0))
    gutter = _find_gutter([(x0, x1) for x0, x1, _ in coords])
    if gutter is None:
        return page_lines
    indices = range(len(page_lines))
    columns = [
        [i for i in indices if coords[i][1] <= gutter],
        [i for i in indices if coords[i][1] > gutter],
    ]
    # The main column (most text) reads first so identity stays at the top and
    # the sidebar's own section headers still split its content afterwards.
    columns.sort(key=lambda col: sum(len(page_lines[i].text) for i in col), reverse=True)
    return [
        page_lines[i]
        for column in columns
        for i in sorted(column, key=lambda idx: (coords[idx][2], coords[idx][0]))
    ]


def _find_gutter(line_intervals: list[tuple[float, float]]) -> float | None:
    """Find an x position crossed by no line, splitting the page in two columns."""
    intervals = sorted(line_intervals)
    best: tuple[float, float] | None = None
    covered_until = intervals[0][1]
    for x0, x1 in intervals[1:]:
        gap = x0 - covered_until
        if gap >= _MIN_GUTTER_WIDTH:
            left = sum(1 for _, line_x1 in line_intervals if line_x1 <= covered_until)
            right = len(line_intervals) - left
            share = min(left, right) / len(line_intervals)
            if (
                min(left, right) >= _MIN_COLUMN_LINES
                and share >= 0.2
                and (best is None or gap > best[0])
            ):
                best = (gap, covered_until)
        covered_until = max(covered_until, x1)
    return best[1] if best else None


def _text_to_document_lines(text: str) -> list[DocumentLine]:
    return [
        DocumentLine(text=line.strip(), line_index=index)
        for index, line in enumerate(text.splitlines())
        if line.strip()
    ]


def _extract_pdf_legacy(data: bytes) -> str:
    from pypdf import PdfReader  # type: ignore[import-not-found]

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def clean_extracted_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
