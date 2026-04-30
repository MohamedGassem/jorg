# backend/services/generation_service.py
"""Generate Word documents by applying candidate profile mappings to a template."""

from __future__ import annotations

import copy
import io
import re
import subprocess
from datetime import date
from pathlib import Path
from typing import Any, Literal
from uuid import UUID

import structlog
from docx import Document  # type: ignore[import-untyped,unused-ignore]
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from core import storage
from core.exceptions import BusinessRuleError, ForbiddenError, NotFoundError
from models.candidate_profile import CandidateProfile, Experience
from models.generated_document import GeneratedDocument
from services import invitation_service, template_service

_PH = re.compile(r"\{\{[^}]+\}\}")

logger = structlog.get_logger()


# ---------- helpers ----------------------------------------------------------


def _fmt_date(d: date | None) -> str:
    return d.strftime("%m/%Y") if d else ""


def _profile_flat(profile: CandidateProfile) -> dict[str, str]:
    return {
        "first_name": profile.first_name or "",
        "last_name": profile.last_name or "",
        "title": profile.title or "",
        "summary": profile.summary or "",
        "phone": profile.phone or "",
        "email_contact": profile.email_contact or "",
        "linkedin_url": profile.linkedin_url or "",
        "location": profile.location or "",
        "years_of_experience": str(profile.years_of_experience or ""),
        "daily_rate": str(profile.daily_rate or ""),
        "annual_salary": str(profile.annual_salary or ""),
        "availability_status": (
            str(profile.availability_status.value) if profile.availability_status else ""
        ),
        "work_mode": str(profile.work_mode.value) if profile.work_mode else "",
        "location_preference": profile.location_preference or "",
        "mission_duration": str(profile.mission_duration.value) if profile.mission_duration else "",
    }


def _exp_flat(exp: Experience) -> dict[str, str]:
    end = _fmt_date(exp.end_date) if not exp.is_current else "présent"
    return {
        "experience.client_name": exp.client_name or "",
        "experience.role": exp.role or "",
        "experience.start_date": _fmt_date(exp.start_date),
        "experience.end_date": end,
        "experience.description": exp.description or "",
        "experience.context": exp.context or "",
        "experience.achievements": exp.achievements or "",
        "experience.technologies": ", ".join(exp.technologies or []),
    }


def _is_text_settable(node: Any) -> bool:
    """Return True if node.text can be assigned (not a read-only computed property)."""
    for klass in type(node).__mro__:
        if "text" in klass.__dict__:
            attr = klass.__dict__["text"]
            if isinstance(attr, property):
                return attr.fset is not None
            # C-level getset_descriptor (lxml native) — always settable
            return True
    return True


def _replace_element(elem: Any, lookup: dict[str, str]) -> None:
    """Replace {{PLACEHOLDER}} in every XML text node using the lookup dict."""
    for node in elem.iter():
        if _is_text_settable(node):
            if node.text:
                node.text = _PH.sub(lambda m: lookup.get(m.group(), ""), node.text)
            if node.tail:
                node.tail = _PH.sub(lambda m: lookup.get(m.group(), ""), node.tail)


def _apply_block(
    doc: Any,
    start_marker: str,
    end_marker: str,
    items: list[dict[str, str]],
    base_lookup: dict[str, str],
) -> None:
    """Clone template paragraphs between markers for each item, then remove markers.

    Uses a while loop to re-scan after each replacement, in case the same
    block appears multiple times in the document.
    """
    while True:
        paras = list(doc.paragraphs)
        start_idx = next((i for i, p in enumerate(paras) if start_marker in p.text), None)
        end_idx = next((i for i, p in enumerate(paras) if end_marker in p.text), None)
        if start_idx is None or end_idx is None:
            break

        # Deep-copy the template XML elements (between markers, exclusive)
        template_elems = [copy.deepcopy(paras[j]._element) for j in range(start_idx + 1, end_idx)]

        anchor = paras[start_idx]._element
        body = doc.element.body

        # Insert clones after anchor (reversed so first item ends up first)
        for item in reversed(items):
            lookup = {**base_lookup, **item}
            for tmpl in reversed(template_elems):
                new_elem = copy.deepcopy(tmpl)
                _replace_element(new_elem, lookup)
                anchor.addnext(new_elem)

        # Remove marker paragraphs and original template paragraphs
        for j in range(start_idx, end_idx + 1):
            body.remove(paras[j]._element)


# ---------- public API -------------------------------------------------------


def generate_document(
    template_path: str,
    profile: CandidateProfile,
    experiences: list[Experience],
    mappings: dict[str, Any],
) -> bytes:
    """Apply mappings to a template docx and return the result as bytes.

    Algorithm:
    1. Build reverse lookup: placeholder → resolved string value.
    2. For experience block markers, clone template paragraphs per experience.
    3. Replace remaining simple placeholders in all paragraphs and table cells.
    4. Return docx bytes.
    """
    doc = Document(template_path)

    profile_data = _profile_flat(profile)

    # Build the simple placeholder → value lookup
    base_lookup: dict[str, str] = {}
    for placeholder, field in mappings.items():
        if not isinstance(field, str):
            continue
        if not field.startswith("experience."):
            base_lookup[placeholder] = profile_data.get(field, "")

    # Build per-experience lookup rows
    exp_items: list[dict[str, str]] = []
    for exp in experiences:
        exp_data = _exp_flat(exp)
        item: dict[str, str] = {}
        for placeholder, field in mappings.items():
            if isinstance(field, str) and field.startswith("experience."):
                item[placeholder] = exp_data.get(field, "")
        exp_items.append(item)

    # Handle {{#EXPERIENCES}}...{{/EXPERIENCES}} blocks
    _apply_block(doc, "{{#EXPERIENCES}}", "{{/EXPERIENCES}}", exp_items, base_lookup)

    # Replace simple placeholders in paragraphs
    for para in doc.paragraphs:
        _replace_element(para._element, base_lookup)

    # Replace simple placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_element(para._element, base_lookup)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert_to_pdf(docx_path: str) -> str | None:
    """Convert a docx file to PDF using LibreOffice headless.

    Returns the PDF path on success, or None if LibreOffice is unavailable
    or conversion fails. Failure is non-fatal: caller delivers docx instead.
    """
    try:
        output_dir = str(Path(docx_path).parent)
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                docx_path,
            ],
            capture_output=True,
            timeout=60,
        )
        if result.returncode == 0:
            pdf_path = str(Path(docx_path).with_suffix(".pdf"))
            if Path(pdf_path).exists():
                return pdf_path
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass
    return None


async def _load_profile(db: AsyncSession, candidate_id: UUID) -> CandidateProfile:
    result = await db.execute(
        select(CandidateProfile).where(CandidateProfile.user_id == candidate_id)
    )
    profile = result.scalar_one_or_none()
    if profile is None:
        raise NotFoundError("candidate_profile_not_found")
    return profile


async def _load_experiences(db: AsyncSession, profile_id: UUID) -> list[Experience]:
    result = await db.execute(select(Experience).where(Experience.profile_id == profile_id))
    return list(result.scalars().all())


async def generate_for_candidate(
    db: AsyncSession,
    organization_id: UUID,
    template_id: UUID,
    candidate_id: UUID,
    generated_by_user_id: UUID,
    fmt: Literal["docx", "pdf"],
) -> GeneratedDocument:
    """Full pipeline: verify grant → load data → generate → save → record."""
    # 1. Verify active AccessGrant
    grant = await invitation_service.get_active_grant(db, candidate_id, organization_id)
    if grant is None:
        raise ForbiddenError("no_active_grant")

    # 2. Verify template is valid and belongs to the organization
    tmpl = await template_service.get_template(db, template_id, organization_id)
    if tmpl is None:
        raise NotFoundError("template_not_found")
    if not tmpl.is_valid:
        raise BusinessRuleError("template_invalid")

    # 3. Load candidate profile and experiences
    profile = await _load_profile(db, candidate_id)
    experiences = await _load_experiences(db, profile.id)

    # 4. Generate docx bytes
    docx_bytes = generate_document(tmpl.word_file_path, profile, experiences, tmpl.mappings)

    # 5. Save docx to local storage
    filename = f"doc_{candidate_id}_{template_id}.docx"
    file_path = storage.save_upload(docx_bytes, filename)

    # 6. Optional PDF conversion (non-fatal on failure)
    actual_path = file_path
    actual_format: str = "docx"
    if fmt == "pdf":
        pdf_path = convert_to_pdf(file_path)
        if pdf_path:
            actual_path = pdf_path
            actual_format = "pdf"

    # 7. Record in DB
    doc = GeneratedDocument(
        access_grant_id=grant.id,
        template_id=template_id,
        generated_by_user_id=generated_by_user_id,
        file_path=actual_path,
        file_format=actual_format,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    logger.info(
        "document.generated",
        template_id=str(template_id),
        candidate_id=str(candidate_id),
        format=fmt,
        access_grant_id=str(doc.access_grant_id),
    )
    return doc


async def list_candidate_documents(db: AsyncSession, candidate_id: UUID) -> list[GeneratedDocument]:
    """All documents generated for a candidate (across all orgs)."""
    from models.invitation import AccessGrant

    result = await db.execute(
        select(GeneratedDocument)
        .join(AccessGrant, GeneratedDocument.access_grant_id == AccessGrant.id)
        .where(AccessGrant.candidate_id == candidate_id)
        .order_by(GeneratedDocument.generated_at.desc())
    )
    return list(result.scalars().all())


async def list_org_documents(db: AsyncSession, organization_id: UUID) -> list[GeneratedDocument]:
    """All documents generated by recruiter's organization."""
    from models.invitation import AccessGrant

    result = await db.execute(
        select(GeneratedDocument)
        .join(AccessGrant, GeneratedDocument.access_grant_id == AccessGrant.id)
        .where(AccessGrant.organization_id == organization_id)
        .order_by(GeneratedDocument.generated_at.desc())
    )
    return list(result.scalars().all())
