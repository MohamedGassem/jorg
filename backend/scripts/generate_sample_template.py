#!/usr/bin/env python3
"""One-shot script to generate backend/static/sample_template.docx."""

from pathlib import Path

from docx import Document


def main() -> None:
    doc = Document()

    doc.add_heading("Dossier de compétences — {{last_name}} {{first_name}}", level=1)

    doc.add_heading("Informations générales", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = table.rows
    rows[0].cells[0].text = "Titre"
    rows[0].cells[1].text = "{{title}}"
    rows[1].cells[0].text = "Localisation"
    rows[1].cells[1].text = "{{location}}"
    rows[2].cells[0].text = "TJM"
    rows[2].cells[1].text = "{{daily_rate}} €/j"
    rows[3].cells[0].text = "Disponibilité"
    rows[3].cells[1].text = "{{availability_status}}"

    doc.add_heading("Résumé", level=2)
    doc.add_paragraph("{{summary}}")

    # Experience block — one paragraph cloned per experience
    doc.add_heading("Expériences professionnelles", level=2)
    doc.add_paragraph("{%p for exp in experiences %}")
    p = doc.add_paragraph()
    run = p.add_run("{{exp.client_name}} — {{exp.role}}")
    run.bold = True
    doc.add_paragraph("Période : {{exp.start_date}} - {{exp.end_date}}")
    doc.add_paragraph("Description : {{exp.description}}")
    doc.add_paragraph("Technologies : {{exp.technologies}}")
    doc.add_paragraph("{%p endfor %}")

    # Skill block — one paragraph cloned per skill
    doc.add_heading("Compétences", level=2)
    doc.add_paragraph("{%p for sk in skills %}")
    doc.add_paragraph("• {{sk.name}} ({{sk.category}})")
    doc.add_paragraph("{%p endfor %}")

    out = Path(__file__).parent.parent / "static" / "sample_template.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Generated: {out}")


if __name__ == "__main__":
    main()
